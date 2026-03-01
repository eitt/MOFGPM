from __future__ import annotations

import argparse
import json
import math
import re
import warnings
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from matplotlib.lines import Line2D


RESOURCE_COLS = ["Doctor", "Nurse", "Assistant", "Specialist"]
DEFAULT_BASELINE_STAFF = {
    "Doctor": 3,
    "Nurse": 3,
    "Assistant": 6,
    "Specialist": 1,
}


def _save_fig(fig, out_path: Path, *, use_tight_layout: bool = True) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if use_tight_layout:
        with warnings.catch_warnings():
            warnings.filterwarnings(
                "ignore",
                message="This figure includes Axes that are not compatible with tight_layout",
                category=UserWarning,
            )
            fig.tight_layout()
    fig.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def _load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _find_latest_run_dir(results_dir: Path) -> Path:
    run_dirs = [p for p in results_dir.glob("run_*") if p.is_dir()]
    if not run_dirs:
        raise FileNotFoundError(f"No run_* folder found in {results_dir}")
    run_dirs = sorted(run_dirs, key=lambda p: p.name)
    return run_dirs[-1]


def _to_numeric(df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
    for c in cols:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")
    return df


def _scenario_order(v) -> int:
    s = str(v)
    m = re.search(r"(\d+)", s)
    if m:
        return int(m.group(1))
    return 10**9


def _load_ga_exports(run_dir: Path, ga_subdir: str) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, dict, float]:
    ga_dir = run_dir / ga_subdir
    results_csv = ga_dir / "results.csv"
    schedule_csv = ga_dir / "schedule_by_sample.csv"
    triage_csv = ga_dir / "triage_waits_by_sample.csv"
    settings_json = ga_dir / "settings.json"

    missing = [p for p in [results_csv, schedule_csv, triage_csv, settings_json] if not p.exists()]
    if missing:
        raise FileNotFoundError(f"Missing GA export files: {missing}")

    results_df = pd.read_csv(results_csv)
    schedule_df = pd.read_csv(schedule_csv)
    triage_df = pd.read_csv(triage_csv)
    settings = _load_json(settings_json)
    horizon = float(settings.get("T_MAX_DAY", 1440.0))

    needed_results = RESOURCE_COLS + ["sample_id", "lambda", "wait", "feasible"]
    miss_res = [c for c in needed_results if c not in results_df.columns]
    if miss_res:
        raise ValueError(f"results.csv missing required columns: {miss_res}")

    needed_schedule = ["sample_id", "patient", "start", "end"]
    miss_sched = [c for c in needed_schedule if c not in schedule_df.columns]
    if miss_sched:
        raise ValueError(f"schedule_by_sample.csv missing required columns: {miss_sched}")

    needed_triage = ["sample_id", "patient", "first_wait"]
    miss_tri = [c for c in needed_triage if c not in triage_df.columns]
    if miss_tri:
        raise ValueError(f"triage_waits_by_sample.csv missing required columns: {miss_tri}")

    results_df = _to_numeric(results_df, RESOURCE_COLS + ["sample_id", "lambda", "wait"])
    schedule_df = _to_numeric(schedule_df, ["sample_id", "start", "end", "task_j"])
    triage_df = _to_numeric(triage_df, ["sample_id", "first_wait"])
    results_df["feasible"] = results_df["feasible"].astype(bool)

    return results_df, schedule_df, triage_df, settings, horizon


def _select_ga_rows_by_staff(results_df: pd.DataFrame) -> pd.DataFrame:
    """
    Step 1 (mandatory):
    if multiple GA rows per staffing, choose feasible=True, then max lambda, then min wait.
    """
    df = results_df.copy()
    df["feasible_priority"] = np.where(df["feasible"], 0, 1)
    sort_cols = RESOURCE_COLS + ["feasible_priority", "lambda", "wait", "sample_id"]
    ascending = [True, True, True, True, True, False, True, True]
    df = df.sort_values(sort_cols, ascending=ascending).reset_index(drop=True)
    counts = df.groupby(RESOURCE_COLS, as_index=False).size().rename(columns={"size": "ga_rows_same_staff"})
    best = df.drop_duplicates(subset=RESOURCE_COLS, keep="first")
    best = best.merge(counts, on=RESOURCE_COLS, how="left")
    return best


def _compute_ga_patient_waits(
    schedule_df: pd.DataFrame,
    triage_df: pd.DataFrame,
    horizon: float,
) -> tuple[pd.DataFrame, int]:
    """
    Step 2 + Step 3:
    - served_24h cohort using last_end <= horizon.
    - patient waiting:
      waiting = (first_start - arrival) + sum(start_current - end_previous)
      where arrival = first_start - first_wait.
    """
    sched = schedule_df.copy()
    sort_cols = ["sample_id", "patient"]
    if "task_j" in sched.columns:
        sort_cols.append("task_j")
    sort_cols.extend(["start", "end"])
    sched = sched.sort_values(sort_cols).reset_index(drop=True)

    patient_bounds = (
        sched.groupby(["sample_id", "patient"], as_index=False)
        .agg(first_start=("start", "min"), last_end=("end", "max"))
    )

    tri = triage_df[["sample_id", "patient", "first_wait"]].drop_duplicates()
    patient_df = patient_bounds.merge(tri, on=["sample_id", "patient"], how="left")
    missing_first_wait = int(patient_df["first_wait"].isna().sum())
    patient_df["first_wait"] = patient_df["first_wait"].fillna(0.0)
    patient_df["arrival_time"] = patient_df["first_start"] - patient_df["first_wait"]

    sched["prev_end"] = sched.groupby(["sample_id", "patient"])["end"].shift(1)
    sched["inter_activity_wait"] = sched["start"] - sched["prev_end"]
    sched["inter_activity_wait"] = sched["inter_activity_wait"].fillna(0.0)
    sched.loc[sched["inter_activity_wait"] < 0.0, "inter_activity_wait"] = 0.0

    internal = (
        sched.groupby(["sample_id", "patient"], as_index=False)["inter_activity_wait"]
        .sum()
        .rename(columns={"inter_activity_wait": "internal_wait_sum"})
    )
    patient_df = patient_df.merge(internal, on=["sample_id", "patient"], how="left")
    patient_df["internal_wait_sum"] = patient_df["internal_wait_sum"].fillna(0.0)
    patient_df["ga_wait_patient"] = patient_df["first_wait"] + patient_df["internal_wait_sum"]
    patient_df["served_24h"] = patient_df["last_end"] <= float(horizon)

    return patient_df, missing_first_wait


def _aggregate_ga_sample_metrics(patient_df: pd.DataFrame) -> pd.DataFrame:
    all_agg = (
        patient_df.groupby("sample_id", as_index=False)
        .agg(
            N_total_ga=("patient", "nunique"),
            N_served_24h_ga=("served_24h", "sum"),
            GA_total_wait_all_finished=("ga_wait_patient", "sum"),
        )
    )

    served_wait = (
        patient_df[patient_df["served_24h"]]
        .groupby("sample_id", as_index=False)["ga_wait_patient"]
        .sum()
        .rename(columns={"ga_wait_patient": "GA_total_wait_served_24h"})
    )
    out = all_agg.merge(served_wait, on="sample_id", how="left")
    out["GA_total_wait_served_24h"] = out["GA_total_wait_served_24h"].fillna(0.0)
    out["GA_avg_wait_all_finished"] = np.where(
        out["N_total_ga"] > 0,
        out["GA_total_wait_all_finished"] / out["N_total_ga"],
        np.nan,
    )
    out["GA_avg_wait_per_served"] = np.where(
        out["N_served_24h_ga"] > 0,
        out["GA_total_wait_served_24h"] / out["N_served_24h_ga"],
        np.nan,
    )
    out["GA_served_rate"] = out["N_served_24h_ga"] / 24.0
    return out


def _detect_replication_columns(df: pd.DataFrame) -> list[str]:
    rep_cols = [c for c in df.columns if re.fullmatch(r"\s*\d+\s*", str(c))]
    if rep_cols:
        return rep_cols
    candidates = []
    for c in df.columns:
        s = str(c).strip().lower()
        if s.startswith("rep") or "replica" in s:
            candidates.append(c)
    return candidates


def _load_flexsim_exports(flexsim_file: Path, sheet_name: str | None = None) -> tuple[pd.DataFrame, list[str], str]:
    if not flexsim_file.exists():
        raise FileNotFoundError(f"FlexSim file not found: {flexsim_file}")

    xls = pd.ExcelFile(flexsim_file)
    if sheet_name:
        if sheet_name not in xls.sheet_names:
            raise ValueError(f"Sheet '{sheet_name}' not found in FlexSim workbook. Available: {xls.sheet_names}")
        sheet = sheet_name
    else:
        sheet = "Raw Data" if "Raw Data" in xls.sheet_names else xls.sheet_names[0]

    fs = pd.read_excel(flexsim_file, sheet_name=sheet)
    required = RESOURCE_COLS + ["Patient"]
    miss = [c for c in required if c not in fs.columns]
    if miss:
        raise ValueError(f"FlexSim sheet '{sheet}' missing required columns: {miss}")

    rep_cols = _detect_replication_columns(fs)
    if not rep_cols:
        raise ValueError("FlexSim export must include replication columns (1..10 or rep*).")

    fs = _to_numeric(fs, RESOURCE_COLS + ["Patient"] + rep_cols)
    fs["FS_mean"] = fs[rep_cols].mean(axis=1)
    fs["FS_std"] = fs[rep_cols].std(axis=1, ddof=1)
    fs["FS_ci95"] = 1.96 * fs["FS_std"] / math.sqrt(float(len(rep_cols)))
    fs["FS_avg_wait_per_served"] = np.where(fs["Patient"] > 0, fs["FS_mean"] / fs["Patient"], np.nan)
    fs["FS_served_rate"] = fs["Patient"] / 24.0
    fs["FS_replications_n"] = int(len(rep_cols))

    keep = ["Scenario"] + RESOURCE_COLS + ["Patient", "FS_mean", "FS_std", "FS_ci95", "FS_avg_wait_per_served", "FS_served_rate", "FS_replications_n"] + rep_cols
    keep = [c for c in keep if c in fs.columns]
    fs = fs[keep].copy()
    fs = fs.drop_duplicates(subset=RESOURCE_COLS, keep="first")
    return fs, rep_cols, sheet


def _extract_baseline_metrics(
    run_dir: Path,
    *,
    ga_subdir: str = "fixed",
    baseline_staff: dict[str, int] | None = None,
    horizon_fallback: float = 1440.0,
) -> tuple[dict, pd.DataFrame]:
    """
    Extract baseline scheduling metrics for a specific staffing tuple to test 24h feasibility.
    """
    staff = dict(DEFAULT_BASELINE_STAFF)
    if baseline_staff:
        staff.update({k: int(v) for k, v in baseline_staff.items() if k in RESOURCE_COLS})

    try:
        res_f, sch_f, tri_f, set_f, horizon_f = _load_ga_exports(run_dir, ga_subdir)
    except Exception:
        return {}, pd.DataFrame()

    horizon = float(set_f.get("T_MAX_DAY", horizon_fallback if np.isfinite(horizon_fallback) else 1440.0))
    target_patients = int(set_f.get("NUM_PATIENTS", 36))

    cand = res_f.copy()
    for r in RESOURCE_COLS:
        cand = cand[pd.to_numeric(cand[r], errors="coerce") == int(staff[r])]

    if cand.empty:
        raise ValueError(
            "Requested baseline staffing tuple not found in GA results: "
            + ", ".join([f"{r}={int(staff[r])}" for r in RESOURCE_COLS])
            + f" (source subdir: {ga_subdir})"
        )

    cand = cand.copy()
    cand["feasible_priority"] = np.where(cand["feasible"], 0, 1)
    cand = cand.sort_values(["feasible_priority", "lambda", "wait", "sample_id"], ascending=[True, False, True, True])
    row = cand.iloc[0]
    sid = int(pd.to_numeric(pd.Series([row.get("sample_id", np.nan)]), errors="coerce").iloc[0])

    p_all, missing_first_wait = _compute_ga_patient_waits(sch_f, tri_f, horizon=horizon)
    p = p_all[p_all["sample_id"] == sid].copy()
    if p.empty:
        raise ValueError(
            f"Baseline sample_id={sid} was found in results.csv but not in schedule/triage exports (source subdir: {ga_subdir})."
        )

    n_total = int(p["patient"].nunique())
    n_served_24h = int((p["served_24h"]).sum())
    completion_max = float(p["last_end"].max())
    served_rate = float(n_served_24h / 24.0)

    out = {
        "sample_id": sid,
        "Doctor": int(row.get("Doctor", np.nan)) if pd.notna(row.get("Doctor", np.nan)) else np.nan,
        "Nurse": int(row.get("Nurse", np.nan)) if pd.notna(row.get("Nurse", np.nan)) else np.nan,
        "Assistant": int(row.get("Assistant", np.nan)) if pd.notna(row.get("Assistant", np.nan)) else np.nan,
        "Specialist": int(row.get("Specialist", np.nan)) if pd.notna(row.get("Specialist", np.nan)) else np.nan,
        "target_patients_24h": target_patients,
        "generated_patients_ga": n_total,
        "served_24h_ga": n_served_24h,
        "served_24h_rate": served_rate,
        "served_share_generated_pct": float(100.0 * n_served_24h / n_total) if n_total > 0 else np.nan,
        "served_share_target_pct": float(100.0 * n_served_24h / target_patients) if target_patients > 0 else np.nan,
        "backlog_vs_generated": int(n_total - n_served_24h),
        "backlog_vs_target": int(max(0, target_patients - n_served_24h)),
        "completion_time_max_min": completion_max,
        "horizon_min": horizon,
        "schedule_feasible_24h_target": bool(n_served_24h >= target_patients and completion_max <= horizon),
        "missing_first_wait_rows": int(missing_first_wait),
        "ga_wait_all_finished": float(row.get("wait", np.nan)),
        "ga_lambda_fixed_row": float(row.get("lambda", np.nan)),
        "ga_feasible_fixed_row": bool(row.get("feasible", False)),
        "baseline_source_subdir": ga_subdir,
    }
    return out, p


def _plot_baseline_completion_profile(
    baseline_patient_df: pd.DataFrame,
    baseline_metrics: dict,
    out_path: Path,
) -> dict[str, str]:
    p = baseline_patient_df.copy()
    p = p.sort_values("last_end", ascending=True).reset_index(drop=True)
    p["idx"] = np.arange(1, len(p) + 1)
    horizon = float(baseline_metrics.get("horizon_min", 1440.0))
    served = p["served_24h"].astype(bool)

    fig, ax = plt.subplots(figsize=(10.5, 5.4))
    ax.scatter(
        p.loc[served, "idx"],
        p.loc[served, "last_end"],
        c="#4C78A8",
        s=35,
        label="Served within 24h",
        edgecolors="#202020",
        linewidths=0.25,
    )
    ax.scatter(
        p.loc[~served, "idx"],
        p.loc[~served, "last_end"],
        c="#E45756",
        s=35,
        label="Completed after 24h",
        edgecolors="#202020",
        linewidths=0.25,
    )
    ax.axhline(horizon, color="#555555", linestyle="--", linewidth=1.2, label=f"24h horizon ({horizon:.0f} min)")
    ax.set_xlabel("Patient index (sorted by completion time)")
    ax.set_ylabel("Patient completion time (minutes)")
    ax.legend(loc="best")
    _save_fig(fig, out_path)

    return {
        "caption": "Figure B1. Baseline completion profile under fixed staffing.",
        "alt_text": "Scatter plot of patient completion times in the fixed baseline case, showing which patients finish within 24 hours and which spill beyond the daily horizon.",
    }


def _plot_patient_heatmap_by_asst_spec(
    df: pd.DataFrame,
    value_col: str,
    out_path: Path,
    *,
    title: str,
    cbar_label: str,
    cmap: str = "viridis",
    vmin: float | None = None,
    vmax: float | None = None,
) -> dict[str, str]:
    assistants = sorted(pd.to_numeric(df["Assistant"], errors="coerce").dropna().astype(int).unique().tolist())
    specialists = sorted(pd.to_numeric(df["Specialist"], errors="coerce").dropna().astype(int).unique().tolist())
    doctors = sorted(pd.to_numeric(df["Doctor"], errors="coerce").dropna().astype(int).unique().tolist())
    nurses = sorted(pd.to_numeric(df["Nurse"], errors="coerce").dropna().astype(int).unique().tolist())

    if not assistants or not specialists or not doctors or not nurses:
        raise ValueError("Insufficient resource levels to build heatmap grid.")

    fig, axes = plt.subplots(
        len(specialists),
        len(assistants),
        figsize=(4.5 * len(assistants) + 1.9, 3.3 * len(specialists) + 1.2),
        squeeze=False,
    )
    ims = []
    for i, sp in enumerate(specialists):
        for j, a in enumerate(assistants):
            ax = axes[i, j]
            sub = df[(df["Assistant"].astype(int) == int(a)) & (df["Specialist"].astype(int) == int(sp))].copy()
            pivot = sub.pivot(index="Nurse", columns="Doctor", values=value_col)
            arr = pivot.reindex(index=nurses, columns=doctors).to_numpy(dtype=float)
            im = ax.imshow(arr, origin="lower", aspect="auto", cmap=cmap, vmin=vmin, vmax=vmax)
            ims.append(im)
            ax.set_xticks(range(len(doctors)))
            ax.set_xticklabels(doctors)
            ax.set_yticks(range(len(nurses)))
            ax.set_yticklabels(nurses)
            if i == len(specialists) - 1:
                ax.set_xlabel("Doctor")
            if j == 0:
                ax.set_ylabel("Nurse")
            for yy in range(arr.shape[0]):
                for xx in range(arr.shape[1]):
                    v = arr[yy, xx]
                    if np.isfinite(v):
                        txt_col = "white" if (vmax is not None and v > (vmin + vmax) / 2.0) else "black"
                        ax.text(xx, yy, f"{v:.1f}", ha="center", va="center", color=txt_col, fontsize=8.5, weight="bold")

    fig.subplots_adjust(left=0.08, right=0.90, bottom=0.08, top=0.93, wspace=0.10, hspace=0.32)
    cax = fig.add_axes([0.915, 0.18, 0.035, 0.66])
    cbar = fig.colorbar(ims[0], cax=cax)
    cbar.set_label(cbar_label)
    _save_fig(fig, out_path, use_tight_layout=False)

    return {
        "caption": title,
        "alt_text": f"Heatmap grid by Assistant and Specialist levels; each panel shows Doctor versus Nurse combinations with annotated values for {cbar_label.lower()}.",
    }


def _build_comparison_dataset(
    fs_df: pd.DataFrame,
    ga_best: pd.DataFrame,
    ga_agg: pd.DataFrame,
    mismatch_threshold_pct: float,
) -> pd.DataFrame:
    ga_cols = RESOURCE_COLS + ["sample_id", "scenario", "lambda", "wait", "feasible", "ga_rows_same_staff"]
    ga_cols = [c for c in ga_cols if c in ga_best.columns]
    ga_sel = ga_best[ga_cols].merge(ga_agg, on="sample_id", how="left")
    ga_sel = ga_sel.rename(
        columns={
            "scenario": "GA_scenario",
            "lambda": "GA_lambda",
            "wait": "GA_wait_results_all_finished",
            "feasible": "GA_feasible",
        }
    )

    merged = fs_df.merge(ga_sel, on=RESOURCE_COLS, how="left")

    merged["GA_scaled_total_24h"] = merged["GA_avg_wait_per_served"] * merged["Patient"]

    merged["diff_per_patient"] = merged["GA_avg_wait_per_served"] - merged["FS_avg_wait_per_served"]
    merged["abs_diff_per_patient"] = merged["diff_per_patient"].abs()
    merged["pct_diff_per_patient"] = np.where(
        merged["FS_avg_wait_per_served"].abs() > 1e-12,
        100.0 * merged["diff_per_patient"] / merged["FS_avg_wait_per_served"],
        np.nan,
    )

    merged["diff_total_24h"] = merged["GA_scaled_total_24h"] - merged["FS_mean"]
    merged["abs_diff_total_24h"] = merged["diff_total_24h"].abs()
    merged["pct_diff_total_24h"] = np.where(
        merged["FS_mean"].abs() > 1e-12,
        100.0 * merged["diff_total_24h"] / merged["FS_mean"],
        np.nan,
    )

    merged["served_gap_patients"] = merged["N_served_24h_ga"] - merged["Patient"]
    merged["served_gap_pct"] = np.where(
        merged["Patient"].abs() > 1e-12,
        100.0 * merged["served_gap_patients"] / merged["Patient"],
        np.nan,
    )
    merged["capacity_horizon_mismatch"] = merged["served_gap_pct"].abs() > float(mismatch_threshold_pct)

    merged["z_score"] = np.where(
        merged["FS_std"].abs() > 1e-12,
        (merged["GA_scaled_total_24h"] - merged["FS_mean"]) / merged["FS_std"],
        np.nan,
    )
    merged["z_interpretation"] = np.where(
        merged["z_score"].isna(),
        "Unverified (FS_std=0)",
        np.where(merged["z_score"].abs() < 2.0, "Consistent (|z|<2)", "Material gap (|z|>=2)"),
    )

    merged["_scenario_order"] = merged["Scenario"].apply(_scenario_order) if "Scenario" in merged.columns else np.arange(len(merged))
    merged = merged.sort_values(["_scenario_order"] + RESOURCE_COLS).reset_index(drop=True)
    return merged


def _build_table1(comp: pd.DataFrame) -> pd.DataFrame:
    cols = [
        "Scenario",
        *RESOURCE_COLS,
        "sample_id",
        "N_total_ga",
        "N_served_24h_ga",
        "GA_served_rate",
        "Patient",
        "FS_served_rate",
        "served_gap_patients",
        "served_gap_pct",
        "capacity_horizon_mismatch",
    ]
    cols = [c for c in cols if c in comp.columns]
    return comp[cols].copy()


def _build_table2(comp: pd.DataFrame) -> pd.DataFrame:
    cols = [
        "Scenario",
        *RESOURCE_COLS,
        "FS_mean",
        "FS_std",
        "FS_ci95",
        "FS_avg_wait_per_served",
        "GA_total_wait_all_finished",
        "GA_avg_wait_all_finished",
        "GA_total_wait_served_24h",
        "GA_avg_wait_per_served",
        "GA_scaled_total_24h",
        "abs_diff_per_patient",
        "pct_diff_per_patient",
        "abs_diff_total_24h",
        "pct_diff_total_24h",
    ]
    cols = [c for c in cols if c in comp.columns]
    return comp[cols].copy()


def _build_table3(comp: pd.DataFrame) -> pd.DataFrame:
    cols = [
        "Scenario",
        *RESOURCE_COLS,
        "GA_scaled_total_24h",
        "FS_mean",
        "FS_std",
        "z_score",
        "z_interpretation",
    ]
    cols = [c for c in cols if c in comp.columns]
    return comp[cols].copy()


def _plot_replication_values_with_ga_reference(df: pd.DataFrame, rep_cols: list[str], out_path: Path) -> dict[str, str]:
    _ = rep_cols  # kept for compatibility; this figure uses aggregated FlexSim statistics.

    d = df.copy()
    d = d[
        d["FS_mean"].notna()
        & d["GA_scaled_total_24h"].notna()
        & np.isfinite(d["FS_mean"])
        & np.isfinite(d["GA_scaled_total_24h"])
        & (d["FS_mean"] > 0)
    ].copy()
    d["ratio_ga_to_fs"] = d["GA_scaled_total_24h"] / d["FS_mean"]

    fig, axes = plt.subplots(1, 2, figsize=(13.0, 5.4))
    ax1, ax2 = axes

    # Panel A: scenario-level ratio view.
    mismatch = d.get("capacity_horizon_mismatch", pd.Series(False, index=d.index)).astype(bool)
    colors = np.where(mismatch, "#E45756", "#4C78A8")
    sizes = 22 + 2.0 * d.get("Patient", pd.Series(0.0, index=d.index)).fillna(0.0).to_numpy(dtype=float)
    ax1.scatter(
        d["FS_mean"],
        d["ratio_ga_to_fs"],
        c=colors,
        s=np.clip(sizes, 22, 90),
        alpha=0.85,
        edgecolors="#202020",
        linewidths=0.3,
    )
    ax1.axhline(1.0, color="#555555", linestyle="--", linewidth=1.2, label="Parity ratio = 1")
    ax1.set_xlabel("FlexSim mean (10 reps)")
    ax1.set_ylabel("GA_scaled_total_24h / FS_mean")
    legend_handles = [
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#4C78A8", markeredgecolor="#202020", markersize=7, label="No capacity/horizon mismatch"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#E45756", markeredgecolor="#202020", markersize=7, label="Capacity/horizon mismatch"),
        Line2D([0], [0], color="#555555", linestyle="--", linewidth=1.2, label="Parity ratio = 1"),
    ]
    ax1.legend(handles=legend_handles, loc="best")

    # Panel B: decile aggregation on ratio + mismatch rate.
    q = int(min(10, max(2, d["FS_mean"].nunique())))
    d["fs_bin"] = pd.qcut(d["FS_mean"], q=q, duplicates="drop")
    grp = (
        d.groupby("fs_bin", as_index=False, observed=False)
        .agg(
            ratio_mean=("ratio_ga_to_fs", "mean"),
            ratio_std=("ratio_ga_to_fs", "std"),
            mismatch_rate=("capacity_horizon_mismatch", "mean"),
            n=("ratio_ga_to_fs", "size"),
        )
    )
    grp["ratio_std"] = grp["ratio_std"].fillna(0.0)
    grp["ratio_ci95"] = 1.96 * grp["ratio_std"] / np.sqrt(grp["n"].clip(lower=1))
    grp["mismatch_rate_pct"] = 100.0 * grp["mismatch_rate"]
    x = np.arange(1, len(grp) + 1)
    ax2_b = ax2.twinx()
    ax2_b.bar(
        x,
        grp["mismatch_rate_pct"].to_numpy(dtype=float),
        color="#E45756",
        alpha=0.18,
        width=0.72,
        label="Mismatch rate (%)",
    )
    ax2.errorbar(
        x,
        grp["ratio_mean"].to_numpy(dtype=float),
        yerr=grp["ratio_ci95"].to_numpy(dtype=float),
        fmt="o-",
        color="#1b4f72",
        ecolor="#5dade2",
        elinewidth=1.1,
        capsize=3,
        label="Mean ratio +/-95% CI (by bin)",
    )
    ax2.axhline(1.0, color="#555555", linestyle="--", linewidth=1.0, label="Parity ratio = 1")
    ax2.set_xlabel("FlexSim mean quantile bin")
    ax2.set_ylabel("GA_scaled_total_24h / FS_mean")
    ax2_b.set_ylabel("Mismatch rate (%)")
    ax2_b.set_ylim(0.0, 100.0)
    ax2.set_xticks(x)
    handles1, labels1 = ax2.get_legend_handles_labels()
    handles2, labels2 = ax2_b.get_legend_handles_labels()
    ax2.legend(handles1 + handles2, labels1 + labels2, loc="best")

    _save_fig(fig, out_path)

    return {
        "caption": "Figure 1. Aggregated ratio view of GA_scaled_total_24h versus FlexSim mean (10 replications).",
        "alt_text": "Left panel shows scenario-level ratio values with blue points for non-mismatch and red points for capacity-horizon mismatch; point size reflects FlexSim served patients. Right panel shows decile-averaged ratio with 95% confidence intervals and mismatch-rate bars.",
    }


def _plot_fs_ci_with_ga_point(df: pd.DataFrame, out_path: Path) -> dict[str, str]:
    n = len(df)
    fig_h = min(max(6.0, n * 0.14), 24.0)
    fig, ax = plt.subplots(figsize=(11.0, fig_h))
    y = np.arange(n)

    ax.errorbar(
        x=df["FS_mean"].to_numpy(dtype=float),
        y=y,
        xerr=df["FS_ci95"].to_numpy(dtype=float),
        fmt="o",
        color="#4C78A8",
        ecolor="#9ecae1",
        elinewidth=1.2,
        capsize=2,
        markersize=4.2,
        label="FlexSim mean +/-95% CI",
    )
    ax.scatter(
        df["GA_scaled_total_24h"].to_numpy(dtype=float),
        y,
        color="tab:red",
        marker="x",
        s=30,
        label="GA_scaled_total_24h",
    )

    ax.set_xlabel("Aggregate KPI value (served-only within 24h)")
    ax.set_ylabel("Scenario index")

    if n <= 35 and "Scenario" in df.columns:
        ax.set_yticks(y)
        ax.set_yticklabels(df["Scenario"].astype(str).tolist())
    else:
        step = max(1, n // 12)
        ticks = np.arange(0, n, step)
        ax.set_yticks(ticks)
        ax.set_yticklabels([str(i + 1) for i in ticks])

    ax.legend(loc="best")
    _save_fig(fig, out_path)

    return {
        "caption": "Figure 2. FlexSim mean with 95% CI and GA scaled 24h point.",
        "alt_text": "Per scenario comparison where blue points and horizontal error bars represent FlexSim mean and confidence interval, and red x markers represent GA scaled 24-hour values.",
    }


def _plot_aggregated_kpi_summary(df: pd.DataFrame, out_path: Path) -> dict[str, str]:
    d = df.copy()
    d = d[
        d["FS_mean"].notna()
        & d["GA_scaled_total_24h"].notna()
        & np.isfinite(d["FS_mean"])
        & np.isfinite(d["GA_scaled_total_24h"])
        & (d["FS_mean"] > 0)
    ].copy()
    d["ratio_ga_to_fs"] = d["GA_scaled_total_24h"] / d["FS_mean"]
    d["mismatch_group"] = np.where(
        d.get("capacity_horizon_mismatch", False).astype(bool),
        "Mismatch",
        "No mismatch",
    )

    def mean_ci95(x: pd.Series) -> tuple[float, float]:
        vals = pd.to_numeric(x, errors="coerce").dropna().to_numpy(dtype=float)
        if vals.size == 0:
            return np.nan, np.nan
        m = float(np.mean(vals))
        if vals.size == 1:
            return m, 0.0
        ci = 1.96 * float(np.std(vals, ddof=1)) / math.sqrt(vals.size)
        return m, ci

    fs_m, fs_ci = mean_ci95(d["FS_mean"])
    ga_m, ga_ci = mean_ci95(d["GA_scaled_total_24h"])

    fig, axes = plt.subplots(1, 2, figsize=(12.8, 5.2))
    ax1, ax2 = axes

    labels = ["FlexSim mean", "GA_scaled_total_24h"]
    vals = [fs_m, ga_m]
    cis = [fs_ci, ga_ci]
    cols = ["#4C78A8", "#E45756"]
    bars = ax1.bar(labels, vals, yerr=cis, color=cols, capsize=5)
    ax1.set_ylabel("Aggregate KPI value")
    for b, v in zip(bars, vals):
        ax1.text(b.get_x() + b.get_width() / 2.0, v, f"{v:.1f}", ha="center", va="bottom", fontsize=9)

    groups = [g for g in ["No mismatch", "Mismatch"] if (d["mismatch_group"] == g).any()]
    data = [d.loc[d["mismatch_group"] == g, "ratio_ga_to_fs"].to_numpy(dtype=float) for g in groups]
    box = ax2.boxplot(data, tick_labels=groups, showfliers=False, patch_artist=True)
    palette = {"No mismatch": "#4C78A8", "Mismatch": "#E45756"}
    for patch, g in zip(box["boxes"], groups):
        patch.set_facecolor(palette[g])
        patch.set_alpha(0.30)
    ax2.axhline(1.0, color="#555555", linestyle="--", linewidth=1.1, label="Parity ratio = 1")
    ax2.set_ylabel("GA_scaled_total_24h / FS_mean")
    ax2.legend(loc="best")

    _save_fig(fig, out_path)
    return {
        "caption": "Figure 2. Aggregated KPI summary: overall levels and ratio distribution.",
        "alt_text": "Left panel compares overall mean aggregate KPI for FlexSim and GA scaled 24-hour values with confidence intervals. Right panel shows the distribution of ratio GA_scaled_total_24h divided by FlexSim mean for mismatch and non-mismatch groups, with parity at 1.",
    }


def _plot_per_patient_bar(df: pd.DataFrame, out_path: Path) -> dict[str, str]:
    ga_mean = float(df["GA_avg_wait_per_served"].mean())
    fs_mean = float(df["FS_avg_wait_per_served"].mean())
    ga_std = float(df["GA_avg_wait_per_served"].std(ddof=1))
    fs_std = float(df["FS_avg_wait_per_served"].std(ddof=1))

    fig, ax = plt.subplots(figsize=(7.0, 4.8))
    labels = ["GA (served-only)", "FlexSim (served-only)"]
    vals = [ga_mean, fs_mean]
    errs = [ga_std, fs_std]
    bars = ax.bar(labels, vals, yerr=errs, color=["#4C78A8", "#F58518"], capsize=4)
    ax.set_ylabel("Per-served-patient KPI")
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width() / 2.0, v, f"{v:.2f}", ha="center", va="bottom", fontsize=10)
    _save_fig(fig, out_path)

    return {
        "caption": "Figure 3. Mean per-served-patient KPI (GA vs FlexSim).",
        "alt_text": "Two-bar chart comparing mean served-only per-patient KPI for GA and FlexSim, with error bars showing scenario variability.",
    }


def _plot_served_rate_bar(df: pd.DataFrame, out_path: Path) -> dict[str, str]:
    ga_mean = float(df["GA_served_rate"].mean())
    fs_mean = float(df["FS_served_rate"].mean())
    ga_std = float(df["GA_served_rate"].std(ddof=1))
    fs_std = float(df["FS_served_rate"].std(ddof=1))

    fig, ax = plt.subplots(figsize=(7.0, 4.8))
    labels = ["GA served rate", "FlexSim served rate"]
    vals = [ga_mean, fs_mean]
    errs = [ga_std, fs_std]
    bars = ax.bar(labels, vals, yerr=errs, color=["#54A24B", "#E45756"], capsize=4)
    ax.set_ylabel("Served patients per hour")
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width() / 2.0, v, f"{v:.2f}", ha="center", va="bottom", fontsize=10)
    _save_fig(fig, out_path)

    return {
        "caption": "Figure 4. Mean served rate per hour (GA vs FlexSim).",
        "alt_text": "Two-bar chart comparing average served patients per hour in GA and FlexSim under the 24-hour served-only cohort definition.",
    }


def _add_word_table(doc: Document, df: pd.DataFrame, caption: str | None = None, max_rows: int = 20) -> None:
    if df.empty:
        p = doc.add_paragraph("[No rows]")
        p.paragraph_format.space_after = Pt(8)
        return

    data = df.head(max_rows).copy()
    t = doc.add_table(rows=1, cols=len(data.columns))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for j, c in enumerate(data.columns):
        hdr[j].text = str(c)

    for _, row in data.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(data.columns):
            v = row[c]
            if isinstance(v, (float, np.floating)):
                cells[j].text = "-" if not np.isfinite(v) else f"{float(v):.4f}"
            else:
                cells[j].text = str(v)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True
        cap.paragraph_format.space_after = Pt(10)


def _write_results_text(
    out_path: Path,
    *,
    n_scenarios: int,
    n_mismatch: int,
    pct_consistent: float,
    ga_total_patients_observed: float,
    ga_served_mean: float,
    fs_served_mean: float,
) -> None:
    lines = [
        "8.1 External validation with FlexSim (served-within-24h reconciliation)",
        "",
        "GA all-finished metrics are not directly comparable to FlexSim 24h outputs.",
        "Validation is performed on the served within 24h cohort to align denominators and avoid the Dylan vs Laura discrepancy.",
        "KPI definition: aggregate waiting/attention metric computed only on patients served within 24h.",
        "FlexSim KPI uses FS_mean (average of 10 replications) and FS_ci95 = 1.96*FS_std/sqrt(10).",
        "GA comparable KPI uses GA_scaled_total_24h = GA_avg_wait_per_served * FlexSim Patient.",
        "Main figures are aggregated for readability; detailed scenario-by-scenario comparison is provided in Appendix Figure A1.",
        "",
        f"Matched staffing scenarios: {n_scenarios}.",
        f"Observed GA total patients per sample (all-finished cohort): {ga_total_patients_observed:.2f}.",
        f"Average GA served within 24h: {ga_served_mean:.2f} patients.",
        f"Average FlexSim served within 24h (Patient): {fs_served_mean:.2f} patients.",
        f"Capacity/horizon mismatch flags: {n_mismatch} scenarios.",
        f"Scenarios with |z|<2 (GA_scaled_total_24h vs FlexSim mean): {pct_consistent:.1f}%.",
        "",
        "Primary comparison uses GA_avg_wait_per_served vs FS_avg_wait_per_served.",
        "Secondary comparison uses GA_scaled_total_24h vs FS_mean.",
    ]
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")


def _write_word_report(
    out_path: Path,
    *,
    run_dir: Path,
    ga_subdir: str,
    flexsim_sheet: str,
    horizon: float,
    missing_first_wait: int,
    mismatch_threshold_pct: float,
    table1: pd.DataFrame,
    table2: pd.DataFrame,
    table3: pd.DataFrame,
    figure_entries: list[dict[str, str]],
    appendix_entries: list[dict[str, str]],
    word_max_rows: int,
) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for h in ["Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[h]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    def add_par(text: str, *, bold: bool = False, align=None, space_after: int = 6):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        if align is not None:
            p.alignment = align
        pf = p.paragraph_format
        pf.space_after = Pt(space_after)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        return p

    p = doc.add_paragraph()
    run = p.add_run("Results Section - FlexSim Validation with 24h Served-Only Cohort")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)

    doc.add_heading("1. Data and Matching Rule", level=1)
    add_par(f"Run folder: {run_dir}")
    add_par(f"GA source subfolder: {ga_subdir}")
    add_par(f"FlexSim sheet used: {flexsim_sheet}")
    add_par(f"24h horizon used: {horizon:.2f} minutes")
    add_par(
        "Matching rule: for each FlexSim staffing row (Doctor, Nurse, Assistant, Specialist), select the GA row with feasible=True, then max lambda, then min wait."
    )

    doc.add_heading("2. Cohort Reconciliation", level=1)
    add_par(
        "GA all-finished metrics are not directly comparable to FlexSim 24h outputs.",
        bold=True,
    )
    add_par(
        "Validation is performed on the served within 24h cohort to align denominators and avoid the Dylan vs Laura discrepancy.",
        bold=True,
    )
    add_par(
        "Served_24h for GA is defined at patient level as last_end <= 1440 (or T_MAX_DAY from settings). "
        "GA waiting on this cohort is computed using first wait plus inter-activity waiting gaps."
    )
    add_par(
        f"Patients with missing first_wait in triage file (filled with 0 for reconstruction): {missing_first_wait}"
    )
    add_par(
        f"Capacity/horizon mismatch threshold used: |served_gap_pct| > {mismatch_threshold_pct:.1f}%"
    )

    doc.add_heading("3. KPI Definition and Figure Encoding", level=1)
    add_par(
        "KPI in this validation refers to the aggregate waiting/attention metric computed only on the served-within-24h cohort. "
        "For FlexSim, FS_mean is the average across 10 replications and FS_ci95 is computed as 1.96*FS_std/sqrt(10)."
    )
    add_par(
        "For GA, GA_scaled_total_24h = GA_avg_wait_per_served * FlexSim Patient, which aligns the GA numerator with the FlexSim served denominator."
    )
    add_par(
        "Figure 1 left panel: each dot is one staffing scenario; blue means no capacity/horizon mismatch, red means mismatch, and larger dots indicate more served patients in FlexSim."
    )
    add_par(
        "Figure 1 right panel: the line and error bars show decile-averaged ratio GA_scaled_total_24h / FS_mean (+/-95% CI across scenarios), and shaded bars show mismatch rate by decile."
    )
    add_par(
        "Figure 2 is an aggregated summary (overall levels and ratio distribution by mismatch group), while Figures 3 and 4 summarize per-patient KPI and served-rate levels."
    )

    n_scenarios = len(table1)
    n_mismatch = int(table1["capacity_horizon_mismatch"].sum()) if "capacity_horizon_mismatch" in table1.columns else 0
    pct_consistent = 100.0 * float((table3["z_interpretation"] == "Consistent (|z|<2)").mean()) if len(table3) > 0 else 0.0
    ga_total_patients_observed = float(table1["N_total_ga"].mean()) if "N_total_ga" in table1.columns else np.nan
    ga_served_mean = float(table1["N_served_24h_ga"].mean()) if "N_served_24h_ga" in table1.columns else np.nan
    fs_served_mean = float(table1["Patient"].mean()) if "Patient" in table1.columns else np.nan

    doc.add_heading("4. Key Summary", level=1)
    add_par(f"Matched scenarios: {n_scenarios}")
    add_par(f"Average GA total patients (all-finished cohort): {ga_total_patients_observed:.2f}")
    add_par(f"Average GA served within 24h: {ga_served_mean:.2f}")
    add_par(f"Average FlexSim served within 24h (Patient): {fs_served_mean:.2f}")
    add_par(f"Capacity/horizon mismatch flags: {n_mismatch}")
    add_par(f"Scenarios with |z|<2: {pct_consistent:.1f}%")

    doc.add_heading("5. Tables", level=1)
    _add_word_table(
        doc,
        table1,
        caption="Table 1. Cohort and capacity reconciliation (full table in CSV).",
        max_rows=word_max_rows,
    )
    _add_word_table(
        doc,
        table2,
        caption="Table 2. Served-only KPI comparison (full table in CSV).",
        max_rows=word_max_rows,
    )
    _add_word_table(
        doc,
        table3,
        caption="Table 3. Statistical check with z-score interpretation (full table in CSV).",
        max_rows=word_max_rows,
    )

    doc.add_heading("6. Figures and Alt Text", level=1)
    for entry in figure_entries:
        fig_path = Path(entry["path"])
        try:
            doc.add_picture(str(fig_path), width=Inches(6.4))
            add_par(entry["caption"], space_after=2)
            add_par(f"Alt text: {entry['alt_text']}", space_after=8)
        except Exception:
            add_par(f"[Unverified] Could not insert figure: {fig_path}")

    if appendix_entries:
        doc.add_heading("Appendix A. Scenario-Level Detail", level=1)
        add_par(
            "Detailed scenario-level comparison is moved to the appendix to keep the main section aggregated and easier to interpret."
        )
        for entry in appendix_entries:
            fig_path = Path(entry["path"])
            try:
                doc.add_picture(str(fig_path), width=Inches(6.4))
                add_par(entry["caption"], space_after=2)
                add_par(f"Alt text: {entry['alt_text']}", space_after=8)
            except Exception:
                add_par(f"[Unverified] Could not insert appendix figure: {fig_path}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def _write_unified_results_text(
    out_path: Path,
    *,
    baseline_metrics: dict,
    n_scenarios: int,
    n_mismatch: int,
    pct_consistent: float,
) -> None:
    lines = [
        "Unified Results Section (Baseline -> Sensitivity -> FlexSim Validation)",
        "",
        "Methodology order follows the baseline-anchored design:",
        "1) baseline feasibility diagnosis, 2) staffing sensitivity analysis, 3) external FlexSim validation.",
        "",
        "Baseline diagnosis:",
        f"- Target patients in 24h: {int(baseline_metrics.get('target_patients_24h', np.nan)) if baseline_metrics else 'Unverified'}",
        f"- Patients generated by GA in baseline sample: {int(baseline_metrics.get('generated_patients_ga', np.nan)) if baseline_metrics else 'Unverified'}",
        f"- Patients served within 24h in baseline sample: {int(baseline_metrics.get('served_24h_ga', np.nan)) if baseline_metrics else 'Unverified'}",
        f"- Maximum completion time in baseline sample (min): {float(baseline_metrics.get('completion_time_max_min', np.nan)):.2f}" if baseline_metrics else "- Maximum completion time: Unverified",
    ]
    if baseline_metrics:
        if bool(baseline_metrics.get("schedule_feasible_24h_target", False)):
            lines.append("- Interpretation: baseline schedule is feasible for the 24h target.")
        else:
            lines.append("- Interpretation: baseline schedule is NOT feasible for the 24h target (backlog remains after 24h).")

    lines.extend(
        [
            "",
            "Sensitivity analysis focuses on average attended patients/day across staffing combinations (Doctor, Nurse, Assistant, Specialist).",
            "Validation then compares GA scaled 24h KPI against FlexSim KPI means from 10 replications.",
            "",
            f"Matched scenarios in validation: {n_scenarios}",
            f"Capacity/horizon mismatch flags: {n_mismatch}",
            f"Scenarios with |z|<2 (GA_scaled_total_24h vs FlexSim mean): {pct_consistent:.1f}%",
        ]
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")


def _write_unified_results_report(
    out_path: Path,
    *,
    run_dir: Path,
    baseline_table: pd.DataFrame,
    sensitivity_table: pd.DataFrame,
    table1: pd.DataFrame,
    table2: pd.DataFrame,
    table3: pd.DataFrame,
    baseline_metrics: dict,
    baseline_figures: list[dict[str, str]],
    sensitivity_figures: list[dict[str, str]],
    validation_figures: list[dict[str, str]],
    appendix_figures: list[dict[str, str]],
    word_max_rows: int,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for h in ["Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[h]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    def add_par(text: str, *, bold: bool = False, align=None, space_after: int = 6):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        if align is not None:
            p.alignment = align
        pf = p.paragraph_format
        pf.space_after = Pt(space_after)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        return p

    p = doc.add_paragraph()
    run = p.add_run("Unified Results Section - Baseline, Sensitivity, and FlexSim Validation")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)

    doc.add_heading("1. Methodology-Aligned Sequence", level=1)
    add_par(f"Run folder: {run_dir}")
    add_par(
        "Results are ordered to match the proposed methodology: (i) baseline scheduling feasibility diagnosis, "
        "(ii) sensitivity analysis with patient throughput by staffing combinations, and (iii) external validation with FlexSim."
    )

    doc.add_heading("2. Baseline Case (24h Scheduling Feasibility)", level=1)
    if baseline_metrics:
        add_par(
            f"Target patients in 24h: {int(baseline_metrics['target_patients_24h'])}; "
            f"GA generated patients: {int(baseline_metrics['generated_patients_ga'])}; "
            f"served within 24h: {int(baseline_metrics['served_24h_ga'])}."
        )
        add_par(
            f"Maximum completion time reached {float(baseline_metrics['completion_time_max_min']):.2f} minutes "
            f"(horizon={float(baseline_metrics['horizon_min']):.0f})."
        )
        if bool(baseline_metrics.get("schedule_feasible_24h_target", False)):
            add_par("Baseline interpretation: feasible for 24h target.", bold=True)
        else:
            add_par("Baseline interpretation: NOT feasible for the 24h target; backlog remains beyond one day.", bold=True)
    _add_word_table(
        doc,
        baseline_table,
        caption="Table B1. Baseline feasibility metrics for fixed staffing.",
        max_rows=word_max_rows,
    )
    for e in baseline_figures:
        try:
            doc.add_picture(str(e["path"]), width=Inches(6.4))
            add_par(e["caption"], space_after=2)
            add_par(f"Alt text: {e['alt_text']}", space_after=8)
        except Exception:
            add_par(f"[Unverified] Could not insert baseline figure: {e['path']}")

    doc.add_heading("3. Sensitivity Analysis on Patient Throughput", level=1)
    add_par(
        "Sensitivity analysis evaluates average attended patients/day under staffing combinations. "
        "Heatmaps are organized by Assistant and Specialist, with Doctor and Nurse axes inside each panel."
    )
    _add_word_table(
        doc,
        sensitivity_table,
        caption="Table S1. Top staffing combinations by GA served-within-24h count.",
        max_rows=word_max_rows,
    )
    for e in sensitivity_figures:
        try:
            doc.add_picture(str(e["path"]), width=Inches(6.4))
            add_par(e["caption"], space_after=2)
            add_par(f"Alt text: {e['alt_text']}", space_after=8)
        except Exception:
            add_par(f"[Unverified] Could not insert sensitivity figure: {e['path']}")

    doc.add_heading("4. External Validation with FlexSim", level=1)
    add_par(
        "Comparability is enforced on the served-within-24h cohort. "
        "KPI definition: aggregate waiting/attention metric for served patients only."
    )
    add_par(
        "FlexSim KPI uses FS_mean (10-replication average). "
        "GA KPI is scaled as GA_scaled_total_24h = GA_avg_wait_per_served * FlexSim Patient."
    )
    _add_word_table(
        doc,
        table1,
        caption="Table V1. Cohort and capacity reconciliation.",
        max_rows=word_max_rows,
    )
    _add_word_table(
        doc,
        table2,
        caption="Table V2. Served-only KPI comparison.",
        max_rows=word_max_rows,
    )
    _add_word_table(
        doc,
        table3,
        caption="Table V3. Statistical check (z-score based).",
        max_rows=word_max_rows,
    )
    for e in validation_figures:
        try:
            doc.add_picture(str(e["path"]), width=Inches(6.4))
            add_par(e["caption"], space_after=2)
            add_par(f"Alt text: {e['alt_text']}", space_after=8)
        except Exception:
            add_par(f"[Unverified] Could not insert validation figure: {e['path']}")

    if appendix_figures:
        doc.add_heading("Appendix A. Detailed Scenario-Level Validation", level=1)
        for e in appendix_figures:
            try:
                doc.add_picture(str(e["path"]), width=Inches(6.4))
                add_par(e["caption"], space_after=2)
                add_par(f"Alt text: {e['alt_text']}", space_after=8)
            except Exception:
                add_par(f"[Unverified] Could not insert appendix figure: {e['path']}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def run_pipeline(args: argparse.Namespace) -> None:
    results_dir = Path(args.results_dir).resolve()
    run_dir = Path(args.run_dir).resolve() if args.run_dir else _find_latest_run_dir(results_dir)
    out_dir = Path(args.outdir).resolve()
    baseline_out_dir = Path(args.baseline_outdir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    baseline_out_dir.mkdir(parents=True, exist_ok=True)
    flexsim_file = Path(args.flexsim_file).resolve()

    results_df, schedule_df, triage_df, settings, horizon = _load_ga_exports(run_dir, args.ga_subdir)
    fs_df, rep_cols, fs_sheet = _load_flexsim_exports(flexsim_file, args.flexsim_sheet)

    ga_best = _select_ga_rows_by_staff(results_df)
    patient_df, missing_first_wait = _compute_ga_patient_waits(schedule_df, triage_df, horizon=horizon)
    ga_agg = _aggregate_ga_sample_metrics(patient_df)

    comp = _build_comparison_dataset(
        fs_df=fs_df,
        ga_best=ga_best,
        ga_agg=ga_agg,
        mismatch_threshold_pct=float(args.mismatch_threshold_pct),
    )

    # Output tables (CSV).
    table1 = _build_table1(comp)
    table2 = _build_table2(comp)
    table3 = _build_table3(comp)
    baseline_staff = {
        "Doctor": int(args.baseline_doctor),
        "Nurse": int(args.baseline_nurse),
        "Assistant": int(args.baseline_assistant),
        "Specialist": int(args.baseline_specialist),
    }
    baseline_metrics, baseline_patient_df = _extract_baseline_metrics(
        run_dir,
        ga_subdir=args.baseline_ga_subdir,
        baseline_staff=baseline_staff,
        horizon_fallback=horizon,
    )
    baseline_table = pd.DataFrame([baseline_metrics]) if baseline_metrics else pd.DataFrame()

    sensitivity_cols = [
        "Doctor",
        "Nurse",
        "Assistant",
        "Specialist",
        "N_served_24h_ga",
        "GA_served_rate",
        "Patient",
        "FS_served_rate",
        "served_gap_patients",
        "capacity_horizon_mismatch",
        "GA_lambda",
    ]
    sensitivity_cols = [c for c in sensitivity_cols if c in comp.columns]
    sensitivity_table = (
        comp.sort_values(["N_served_24h_ga", "GA_lambda", "Patient"], ascending=[False, False, False])[sensitivity_cols]
        .head(15)
        .reset_index(drop=True)
    )

    p_table1 = out_dir / "table1_cohort_capacity_reconciliation.csv"
    p_table2 = out_dir / "table2_kpi_comparison_served_only.csv"
    p_table3 = out_dir / "table3_statistical_check.csv"
    p_dataset = out_dir / "validation_matched_dataset_24h.csv"
    p_patient = out_dir / "ga_patient_level_waits_24h.csv"
    p_baseline_table = out_dir / "baseline_fixed_case_summary.csv"
    p_baseline_table_dedicated = baseline_out_dir / "baseline_fixed_case_summary.csv"
    p_sensitivity_table = out_dir / "sensitivity_top_staffing_by_served24h.csv"

    table1.to_csv(p_table1, index=False)
    table2.to_csv(p_table2, index=False)
    table3.to_csv(p_table3, index=False)
    comp.to_csv(p_dataset, index=False)
    patient_df.to_csv(p_patient, index=False)
    baseline_table.to_csv(p_baseline_table, index=False)
    baseline_table.to_csv(p_baseline_table_dedicated, index=False)
    sensitivity_table.to_csv(p_sensitivity_table, index=False)

    # Figures.
    figure_entries: list[dict[str, str]] = []
    baseline_entries: list[dict[str, str]] = []
    sensitivity_entries: list[dict[str, str]] = []

    f1 = out_dir / "fig1_aggregated_ratio_view.png"
    meta1 = _plot_replication_values_with_ga_reference(comp, rep_cols, f1)
    figure_entries.append({"path": str(f1), **meta1})

    f2 = out_dir / "fig2_aggregated_kpi_summary.png"
    meta2 = _plot_aggregated_kpi_summary(comp, f2)
    figure_entries.append({"path": str(f2), **meta2})

    f3 = out_dir / "fig3_per_served_patient_kpi_bar.png"
    meta3 = _plot_per_patient_bar(comp, f3)
    figure_entries.append({"path": str(f3), **meta3})

    f4 = out_dir / "fig4_served_rate_per_hour_bar.png"
    meta4 = _plot_served_rate_bar(comp, f4)
    figure_entries.append({"path": str(f4), **meta4})

    appendix_entries: list[dict[str, str]] = []
    fA1 = out_dir / "appendix_figA1_scenario_ci_vs_ga_scaled.png"
    metaA1 = _plot_fs_ci_with_ga_point(comp, fA1)
    metaA1["caption"] = "Appendix Figure A1. Scenario-level FlexSim mean +/-95% CI with GA scaled 24h points."
    metaA1["alt_text"] = "Detailed scenario-by-scenario chart where each row shows FlexSim mean with confidence interval and a GA scaled 24-hour marker."
    appendix_entries.append({"path": str(fA1), **metaA1})

    p_baseline_figure_dedicated = baseline_out_dir / "figB1_baseline_completion_profile.png"
    if not baseline_patient_df.empty and baseline_metrics:
        fB1 = out_dir / "figB1_baseline_completion_profile.png"
        mB1 = _plot_baseline_completion_profile(baseline_patient_df, baseline_metrics, fB1)
        baseline_entries.append({"path": str(fB1), **mB1})
        if p_baseline_figure_dedicated.resolve() != fB1.resolve():
            _plot_baseline_completion_profile(baseline_patient_df, baseline_metrics, p_baseline_figure_dedicated)

    # Sensitivity figures focused on attended-patient patterns by staffing combinations.
    fS1 = out_dir / "figS1_ga_served24h_by_staff_grid.png"
    mS1 = _plot_patient_heatmap_by_asst_spec(
        comp,
        "N_served_24h_ga",
        fS1,
        title="Figure S1. GA served patients within 24h by staffing combination",
        cbar_label="GA served patients within 24h",
        cmap="viridis",
    )
    sensitivity_entries.append({"path": str(fS1), **mS1})

    fS2 = out_dir / "figS2_flexsim_patient_by_staff_grid.png"
    mS2 = _plot_patient_heatmap_by_asst_spec(
        comp,
        "Patient",
        fS2,
        title="Figure S2. FlexSim average attended patients/day by staffing combination",
        cbar_label="FlexSim avg attended patients/day",
        cmap="viridis",
    )
    sensitivity_entries.append({"path": str(fS2), **mS2})

    comp["served_gap_fs_minus_ga"] = comp["Patient"] - comp["N_served_24h_ga"]
    max_abs_gap = float(np.nanmax(np.abs(pd.to_numeric(comp["served_gap_fs_minus_ga"], errors="coerce"))))
    fS3 = out_dir / "figS3_gap_flexsim_minus_ga_by_staff_grid.png"
    mS3 = _plot_patient_heatmap_by_asst_spec(
        comp,
        "served_gap_fs_minus_ga",
        fS3,
        title="Figure S3. Gap in attended patients/day (FlexSim - GA served_24h)",
        cbar_label="Patient gap (FlexSim - GA)",
        cmap="coolwarm",
        vmin=-max_abs_gap,
        vmax=max_abs_gap,
    )
    sensitivity_entries.append({"path": str(fS3), **mS3})

    # Alt text output.
    p_alt = out_dir / "alt_text_figures.txt"
    alt_lines = []
    for e in baseline_entries + sensitivity_entries + figure_entries + appendix_entries:
        alt_lines.append(f"Figure: {Path(e['path']).name}")
        alt_lines.append(f"Caption: {e['caption']}")
        alt_lines.append(f"Alt text: {e['alt_text']}")
        alt_lines.append("")
    p_alt.write_text("\n".join(alt_lines).strip() + "\n", encoding="utf-8")

    # Results section text + Word.
    p_text = out_dir / "Results_Validation_FlexSim_24h.txt"
    n_scenarios = len(table1)
    n_mismatch = int(table1["capacity_horizon_mismatch"].sum()) if len(table1) > 0 else 0
    pct_consistent = 100.0 * float((table3["z_interpretation"] == "Consistent (|z|<2)").mean()) if len(table3) > 0 else 0.0
    ga_total_patients_observed = float(table1["N_total_ga"].mean()) if len(table1) > 0 else np.nan
    ga_served_mean = float(table1["N_served_24h_ga"].mean()) if len(table1) > 0 else np.nan
    fs_served_mean = float(table1["Patient"].mean()) if len(table1) > 0 else np.nan
    _write_results_text(
        p_text,
        n_scenarios=n_scenarios,
        n_mismatch=n_mismatch,
        pct_consistent=pct_consistent,
        ga_total_patients_observed=ga_total_patients_observed,
        ga_served_mean=ga_served_mean,
        fs_served_mean=fs_served_mean,
    )

    p_word = out_dir / "Results_Validation_FlexSim_24h.docx"
    _write_word_report(
        p_word,
        run_dir=run_dir,
        ga_subdir=args.ga_subdir,
        flexsim_sheet=fs_sheet,
        horizon=horizon,
        missing_first_wait=missing_first_wait,
        mismatch_threshold_pct=float(args.mismatch_threshold_pct),
        table1=table1,
        table2=table2,
        table3=table3,
        figure_entries=figure_entries,
        appendix_entries=appendix_entries,
        word_max_rows=int(args.word_max_rows),
    )

    p_unified_text = out_dir / "Unified_Results_Section_MOFGPM_FlexSim.txt"
    _write_unified_results_text(
        p_unified_text,
        baseline_metrics=baseline_metrics,
        n_scenarios=n_scenarios,
        n_mismatch=n_mismatch,
        pct_consistent=pct_consistent,
    )

    p_unified_word = out_dir / "Unified_Results_Section_MOFGPM_FlexSim.docx"
    _write_unified_results_report(
        p_unified_word,
        run_dir=run_dir,
        baseline_table=baseline_table,
        sensitivity_table=sensitivity_table,
        table1=table1,
        table2=table2,
        table3=table3,
        baseline_metrics=baseline_metrics,
        baseline_figures=baseline_entries,
        sensitivity_figures=sensitivity_entries,
        validation_figures=figure_entries,
        appendix_figures=appendix_entries,
        word_max_rows=int(args.word_max_rows),
    )

    # Also provide a concise execution summary.
    p_summary = out_dir / "validation_summary.json"
    summary = {
        "run_dir": str(run_dir),
        "ga_subdir": args.ga_subdir,
        "baseline_ga_subdir": args.baseline_ga_subdir,
        "baseline_staff": baseline_staff,
        "baseline_outdir": str(baseline_out_dir),
        "flexsim_file": str(flexsim_file),
        "flexsim_sheet": fs_sheet,
        "horizon_minutes": float(horizon),
        "configured_num_patients": settings.get("NUM_PATIENTS"),
        "matched_scenarios": int(n_scenarios),
        "mean_ga_total_patients_all_finished": float(ga_total_patients_observed),
        "mean_ga_served_24h": float(ga_served_mean),
        "mean_fs_served_24h": float(fs_served_mean),
        "capacity_horizon_mismatch_count": int(n_mismatch),
        "z_consistent_pct": float(pct_consistent),
        "baseline_target_patients_24h": baseline_metrics.get("target_patients_24h") if baseline_metrics else None,
        "baseline_generated_patients_ga": baseline_metrics.get("generated_patients_ga") if baseline_metrics else None,
        "baseline_served_24h_ga": baseline_metrics.get("served_24h_ga") if baseline_metrics else None,
        "baseline_schedule_feasible_24h_target": baseline_metrics.get("schedule_feasible_24h_target") if baseline_metrics else None,
    }
    p_summary.write_text(json.dumps(summary, indent=2), encoding="utf-8")

    print(f"Run folder used: {run_dir}")
    print(f"GA source folder: {run_dir / args.ga_subdir}")
    print(f"Baseline GA source folder: {run_dir / args.baseline_ga_subdir}")
    print("Baseline staffing tuple: " + ", ".join([f"{k}={v}" for k, v in baseline_staff.items()]))
    print(f"Baseline dedicated output folder: {baseline_out_dir}")
    print(f"FlexSim file used: {flexsim_file} (sheet: {fs_sheet})")
    print(f"Saved: {p_table1}")
    print(f"Saved: {p_table2}")
    print(f"Saved: {p_table3}")
    print(f"Saved: {p_dataset}")
    print(f"Saved: {p_patient}")
    print(f"Saved: {p_baseline_table}")
    print(f"Saved: {p_baseline_table_dedicated}")
    print(f"Saved: {p_sensitivity_table}")
    for e in baseline_entries:
        print(f"Saved: {e['path']}")
    if p_baseline_figure_dedicated.exists():
        print(f"Saved: {p_baseline_figure_dedicated}")
    for e in sensitivity_entries:
        print(f"Saved: {e['path']}")
    for e in figure_entries:
        print(f"Saved: {e['path']}")
    for e in appendix_entries:
        print(f"Saved: {e['path']}")
    print(f"Saved: {p_alt}")
    print(f"Saved: {p_text}")
    print(f"Saved: {p_word}")
    print(f"Saved: {p_unified_text}")
    print(f"Saved: {p_unified_word}")
    print(f"Saved: {p_summary}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Transform existing GA and FlexSim exports to validate with a comparable 24h served-only cohort. "
            "This script does not rerun GA or FlexSim."
        )
    )
    parser.add_argument("--results-dir", default="results", help="Base results directory.")
    parser.add_argument(
        "--run-dir",
        default=None,
        help="Specific run folder (e.g., results/run_20260228_091617). If omitted, latest run_* under --results-dir.",
    )
    parser.add_argument(
        "--ga-subdir",
        default="combinatorics",
        help="Subfolder inside run dir that contains results.csv/schedule_by_sample.csv/triage_waits_by_sample.csv/settings.json.",
    )
    parser.add_argument("--flexsim-file", default="results/Flexsim/10_replicas.xlsx", help="FlexSim workbook path.")
    parser.add_argument(
        "--flexsim-sheet",
        default=None,
        help="Optional FlexSim sheet name. If omitted: 'Raw Data' when available, else first sheet.",
    )
    parser.add_argument(
        "--outdir",
        default="results/validation_flexsim_24h",
        help="Output directory for tables, figures, alt text, and results section.",
    )
    parser.add_argument(
        "--baseline-outdir",
        default="results/result_baseline",
        help="Dedicated output directory for baseline-only artifacts.",
    )
    parser.add_argument(
        "--baseline-ga-subdir",
        default="fixed",
        help="GA subfolder used for baseline extraction (inside run_dir).",
    )
    parser.add_argument(
        "--baseline-doctor",
        type=int,
        default=DEFAULT_BASELINE_STAFF["Doctor"],
        help="Baseline Doctor count used to select baseline scenario.",
    )
    parser.add_argument(
        "--baseline-nurse",
        type=int,
        default=DEFAULT_BASELINE_STAFF["Nurse"],
        help="Baseline Nurse count used to select baseline scenario.",
    )
    parser.add_argument(
        "--baseline-assistant",
        type=int,
        default=DEFAULT_BASELINE_STAFF["Assistant"],
        help="Baseline Assistant count used to select baseline scenario.",
    )
    parser.add_argument(
        "--baseline-specialist",
        type=int,
        default=DEFAULT_BASELINE_STAFF["Specialist"],
        help="Baseline Specialist count used to select baseline scenario.",
    )
    parser.add_argument(
        "--mismatch-threshold-pct",
        type=float,
        default=20.0,
        help="Threshold for capacity/horizon mismatch flag using |(GA_served - FlexSim_served)/FlexSim_served|*100.",
    )
    parser.add_argument(
        "--word-max-rows",
        type=int,
        default=20,
        help="Maximum rows per table inserted in the Word report.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    run_pipeline(args)
