from pathlib import Path
import re
import shutil
import subprocess

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# ============================================================
# WORD DOC BUILDER — aligned to the enhanced LaTeX (with prose)
# Produces a Methodology section that:
# - explains what each equation *means* (managerial interpretation)
# - includes Table 1 (staff/cost/bounds) and Table 2 (fuzzy mapping)
# - includes Jiménez linearization and EI defuzzification
# - includes triage feasibility (early-care SLA) equation
# - includes MOFGPM membership functions and max–min model
# ============================================================

doc = Document()

# ------------------
# Page setup / styles
# ------------------
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

for h in ["Heading 1", "Heading 2", "Heading 3"]:
    st = doc.styles[h]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_par(
    text,
    bold=False,
    italic=False,
    align=None,
    space_after=6,
    space_before=0,
    line_spacing=1.15,
):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    return p


# ------------------
# LaTeX-like -> Word equation safe text
# (lightweight; good for linear math in OMML m:t)
# ------------------
def latex_to_word_math_text(expr: str) -> str:
    text = expr

    def _extract_braced(s, start_idx):
        if start_idx >= len(s) or s[start_idx] != "{":
            return None, start_idx
        depth = 0
        content_start = start_idx + 1
        for i in range(start_idx, len(s)):
            if s[i] == "{":
                depth += 1
            elif s[i] == "}":
                depth -= 1
                if depth == 0:
                    return s[content_start:i], i + 1
        return None, start_idx

    def _replace_frac(s):
        i = 0
        out = []
        while i < len(s):
            if s.startswith(r"\frac", i):
                j = i + len(r"\frac")
                while j < len(s) and s[j].isspace():
                    j += 1
                num, j_after_num = _extract_braced(s, j)
                if num is None:
                    out.append(s[i])
                    i += 1
                    continue
                j = j_after_num
                while j < len(s) and s[j].isspace():
                    j += 1
                den, j_after_den = _extract_braced(s, j)
                if den is None:
                    out.append(s[i])
                    i += 1
                    continue
                out.append(f"({_replace_frac(num)})/({_replace_frac(den)})")
                i = j_after_den
            else:
                out.append(s[i])
                i += 1
        return "".join(out)

    text = _replace_frac(text)

    # accents: \tilde{T} -> T̃
    text = re.sub(r"\\tilde\{([^{}]+)\}", lambda m: f"{m.group(1)}\u0303", text)

    replacements = {
        r"\sum": "Σ",
        r"\cdot": "·",
        r"\sigma": "σ",
        r"\mu": "μ",
        r"\lambda": "λ",
        r"\ge": "≥",
        r"\le": "≤",
        r"\forall": "∀",
        r"\max": "max",
        r"\quad": " ",
        r"\ ": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)

    # subscripts/superscripts groups: x_{ij} -> x_ij ; x^{2} -> x^2
    text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)

    # remove braces/backslashes
    text = text.replace("{", "").replace("}", "")
    text = text.replace("\\", "")

    text = re.sub(r"\s+", " ", text).strip()
    return text


def add_word_equation(eq_text: str, eq_label: str | None = None):
    """Insert a Word equation object (OMML) with pre-normalized text."""
    if eq_label:
        add_par(eq_label, bold=True, space_after=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = eq_text
    mr.append(mt)
    omath.append(mr)
    omath_para.append(omath)
    p._p.append(omath_para)


def add_equation_block(lines: list[str], eq_label: str | None = None):
    if eq_label:
        add_par(eq_label, bold=True, space_after=2)
    for ln in lines:
        add_word_equation(latex_to_word_math_text(ln))
    add_par("", space_after=8)


def add_table_caption(text: str):
    add_par(text, italic=True, space_before=2, space_after=10)


# ============================================================
# TITLE + INTRO
# ============================================================
title = doc.add_paragraph()
title_run = title.add_run(
    "Methodology — Baseline-Anchored Sensitivity Analysis with MOFGPM and Jiménez Linearization"
)
title_run.bold = True
title_run.font.size = Pt(14)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)

add_par(
    "This section presents a decision-support pipeline for Emergency Department (ED) resource planning under uncertainty. "
    "The methodology is intentionally baseline-anchored: it begins from a realistic staffing configuration already used (or constrained) "
    "in practice, and then extends the analysis through sensitivity runs to (i) diagnose feasibility of service targets under uncertainty "
    "and (ii) map interpretable tradeoffs between service satisfaction and cost. Rather than producing a single point recommendation, the "
    "pipeline produces a feasibility envelope and a cost–service tradeoff surface that can support managerial negotiation."
)

# ============================================================
# 1. PROBLEM STATEMENT
# ============================================================
doc.add_heading("1. Problem statement and baseline-anchored design", level=1)

add_par(
    "Patients arrive over time, are assigned a triage category, and follow triage-dependent activity routes. "
    "Each activity consumes one or more shared resources (Doctor, Nurse, Assistant, Specialist). "
    "A staffing plan specifies the available capacity per resource type."
)

add_par(
    "Given a staffing plan, the pipeline constructs a feasible schedule (respecting patient precedence and resource availability), "
    "computes waiting-time performance and staffing cost, and evaluates goal satisfaction using max–min Multiobjective Fuzzy Goal Programming (MOFGPM). "
    "The baseline-anchored design addresses the practical decision context: managers start from an existing staffing plan and need both a feasibility "
    "diagnosis and an interpretable tradeoff map rather than a one-shot optimum."
)

# ============================================================
# 2. CASE INPUTS (TABLE 1) + COST EQ WITH INTERPRETATION
# ============================================================
doc.add_heading("2. Case inputs: resources, staffing, and cost", level=1)

add_par(
    "Staffing is analyzed in two modes: (i) a baseline fixed staffing configuration representing current operations or policy constraints, "
    "and (ii) sensitivity sampling within lower/upper bounds to explore alternative capacity levels. "
    "Table 1 summarizes the resource types, unit costs, baseline staffing, and practical bounds used in sensitivity runs."
)

table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Resource type"
hdr[1].text = "Unit cost (per staff unit)"
hdr[2].text = "Baseline staff (example)"
hdr[3].text = "Lower bound (sensitivity)"
hdr[4].text = "Upper bound (sensitivity)"

resources = [
    ("Doctor", 200, 3, 1, 6),
    ("Nurse", 100, 3, 1, 8),
    ("Assistant", 50, 6, 1, 10),
    ("Specialist", 300, 1, 0, 2),
]
for r, c, b, lo, hi in resources:
    row = table.add_row().cells
    row[0].text = str(r)
    row[1].text = str(c)
    row[2].text = str(b)
    row[3].text = str(lo)
    row[4].text = str(hi)

add_table_caption("Table 1. Resource types, unit costs, baseline staffing, and sensitivity bounds.")

add_par(
    "For a staffing vector S = {S_r}, total staffing cost is computed as the linear aggregation of unit costs and staffing quantities:"
)
add_equation_block(
    [r"\mathrm{Cost}(S) = \sum_{r} SC_{r}\,S_{r}"],
    eq_label="Equation (1). Staffing cost"
)

add_par(
    "Equation (1) connects capacity decisions to financial sustainability. Increasing staffing increases the available capacity to meet demand "
    "and reduce waiting, but it also increases cost. This cost term later enters MOFGPM through a satisfaction (membership) function that reflects "
    "a budget aspiration (goal) and an unacceptable upper limit."
)

# ============================================================
# 3. UNCERTAINTY MODELING
# ============================================================
doc.add_heading("3. Uncertainty modeling", level=1)

doc.add_heading("3.1 Patient arrivals by triage (demand model)", level=2)
add_par(
    "Patient arrivals are generated by triage category using fitted probabilistic models. Inter-arrival times are simulated and accumulated to obtain "
    "arrival timestamps for a representative day. Using the same realized arrival stream across uncertainty scenarios (when desired) isolates the effect "
    "of service-time uncertainty and staffing on feasibility and performance."
)

add_par("Notation for arrivals:")
add_equation_block(
    [r"TA_i:\ \text{arrival time of patient } i,\quad c(i):\ \text{triage class of patient } i"],
    eq_label="Notation (arrivals)"
)
add_par(
    "Arrival times {TA_i} define the time-varying demand pressure. When arrivals overlap heavily, a fixed staffing plan may become infeasible with respect "
    "to triage service standards, even if the same plan is feasible on a low-demand day."
)

doc.add_heading("3.2 Service-time uncertainty and fuzzy mapping (Jiménez + scenarios)", level=2)
add_par(
    "Service durations are uncertain due to heterogeneity in patient condition, interruptions, and clinical complexity. To integrate heterogeneous uncertainty "
    "descriptions (uniform, triangular, exponential) within one framework, each activity duration is mapped to a fuzzy triple (p, m, o) representing an "
    "optimistic value p, a central value m, and a pessimistic value o. This supports scenario-based analysis while retaining a single defuzzification rule."
)

add_par(
    "For triangular activities, the fuzzy triple is defined directly as (p, m, o) = (L, M, U), where M is the most likely duration. "
    "For uniform activities over [a, b], the triple is (p, m, o) = (a, (a+b)/2, b). "
    "For exponential activities, we retain Jiménez’s linearization to obtain a bounded representation from a location parameter a and scale σ, "
    "then define a central value a + σ/2 to complete the fuzzy triple."
)

add_equation_block(
    [
        r"(p,m,o)=(L,M,U)\ \ \text{for triangular}",
        r"(p,m,o)=\left(a,\frac{a+b}{2},b\right)\ \ \text{for uniform }(a,b)",
        r"(p,m,o)=\left(a,\ a+\frac{\sigma}{2},\ a+\sigma\right)\ \ \text{for exponential via Jimenez}",
    ],
    eq_label="Equations (2a–2c). Fuzzy mapping (including Jiménez linearization)"
)

add_par(
    "Equations (2a–2c) define how uncertainty in durations is represented in a consistent three-point form. In particular, the Jiménez mapping "
    "preserves the interval concept [a, a+σ] while providing a central value a+σ/2 for computational use. In implementation, the location parameter a "
    "can be set to the mean or median of the fitted exponential distribution, depending on robustness preferences."
)

add_par(
    "For the expected (defuzzified) scenario used in schedule evaluation, the pipeline applies an expected-value defuzzification operator:"
)
add_equation_block(
    [r"EI(\tilde{T}) = \frac{p + 2m + o}{4}"],
    eq_label="Equation (3). Defuzzification (expected-value)"
)
add_par(
    "Equation (3) balances optimistic and pessimistic extremes while placing higher weight on the central value. This yields a single crisp duration "
    "used by the scheduling-based evaluation module."
)

# Table 2 (mapping summary)
add_par(
    "Table 2 summarizes how original uncertainty specifications are mapped into fuzzy triples and how the three scenario values are defined "
    "(optimistic, expected, pessimistic)."
)

table2 = doc.add_table(rows=1, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
h = table2.rows[0].cells
h[0].text = "Original specification"
h[1].text = "Fuzzy triple (p, m, o)"
h[2].text = "Expected scenario value"
h[3].text = "Scenario values used"

rows2 = [
    ("Uniform(a,b)", "(a, (a+b)/2, b)", "EI = (p+2m+o)/4", "optimistic=p; expected=EI; pessimistic=o"),
    ("Triangular(L,M,U)", "(L, M, U)", "EI = (L+2M+U)/4", "optimistic=L; expected=EI; pessimistic=U"),
    ("Exponential(a,σ) via Jiménez", "(a, a+σ/2, a+σ)", "EI = (p+2m+o)/4", "optimistic=p; expected=EI; pessimistic=o"),
]
for a, b, c, d in rows2:
    r = table2.add_row().cells
    r[0].text = a
    r[1].text = b
    r[2].text = c
    r[3].text = d

add_table_caption(
    "Table 2. Mapping from probabilistic/triangular inputs to fuzzy scenarios and defuzzified expected values."
)

# ============================================================
# 4. SCHEDULING-BASED EVALUATION (DEMAND vs CAPACITY)
# ============================================================
doc.add_heading("4. Scheduling-based evaluation (capacity meets demand)", level=1)

add_par(
    "Given arrivals (demand) and staffing (capacity), the pipeline constructs a feasible schedule for all required patient activities. "
    "Let ST_{ij} and ET_{ij} denote the start and end times of activity j for patient i, respectively, TA_i the arrival time, and d_{ij} "
    "the defuzzified duration (typically d_{ij} = EI(\\tilde{T}_{ij}) from Equation (3)). Precedence and timing are:"
)

add_equation_block(
    [
        r"ST_{i1} \ge TA_{i}",
        r"ST_{ij} \ge ET_{i,j-1},\quad \forall j>1",
        r"ET_{ij} = ST_{ij} + d_{ij}",
    ],
    eq_label="Equations (4–6). Timing and precedence"
)

add_par(
    "Equations (4–6) enforce temporal feasibility. Equation (4) prevents service before arrival. Equation (5) enforces within-patient precedence "
    "(the clinical route must be followed). Equation (6) links start times to completion times using uncertainty-adjusted durations. "
    "Operationally, these constraints define how demand propagates through the ED over time."
)

add_par(
    "Resource feasibility (capacity) is enforced through a serial schedule generation scheme: when an activity requires one or more resource types, "
    "it is assigned to the earliest available units of each required resource type, and its start time is set to the latest of (i) the precedence-implied "
    "release time and (ii) the selected resource availability times. This operationally encodes the balance between demand and capacity: if arrivals overlap "
    "and staffing is limited, resource availability delays start times and increases waiting."
)

add_par("Waiting time is computed as:")
add_equation_block(
    [
        r"WT_{i1} = ST_{i1} - TA_{i}",
        r"WT_{ij} = ST_{ij} - ET_{i,j-1},\quad \forall j>1",
        r"\mathrm{Wait}(S) = \sum_{i}\sum_{j} WT_{ij}",
    ],
    eq_label="Equations (7–9). Waiting-time metric"
)

add_par(
    "Equations (7–9) define the service-quality metric used in MOFGPM. The aggregate Wait(S) is the total accumulated waiting time across all patients and "
    "activities under staffing plan S. This quantity typically decreases when staffing capacity increases and increases when demand intensity or service durations "
    "increase."
)

# Triage feasibility rule (added; present in enhanced LaTeX)
add_par(
    "In addition to aggregate waiting, feasibility is evaluated against triage-specific service standards. Let WT_i^{early} denote the cumulative waiting time "
    "over early-care activities (e.g., up to a fixed index j ≤ J0 in the implementation), and let MT_c be the maximum tolerated threshold for triage category c. "
    "Feasibility requires:"
)
add_equation_block(
    [r"WT^{early}_{i} \le MT_{c(i)},\quad \forall i"],
    eq_label="Equation (10). Triage feasibility (early-care service standard)"
)
add_par(
    "Equation (10) ensures that high-acuity patients are not delayed beyond clinically acceptable thresholds, even if aggregate waiting is acceptable. "
    "In computation, schedules violating this constraint are penalized to prioritize feasibility before improving satisfaction."
)

# ============================================================
# 5. MOFGPM
# ============================================================
doc.add_heading("5. Multiobjective Fuzzy Goal Programming (MOFGPM)", level=1)

add_par(
    "Cost and waiting are conflicting objectives: increasing staffing reduces waiting but increases cost. Instead of applying fixed weights, MOFGPM converts each "
    "objective into a satisfaction degree in [0,1] using piecewise-linear membership functions defined by aspiration (Goal) and unacceptable (Max) levels."
)

add_par("Cost membership function:")
add_equation_block(
    [
        r"\mu_{\mathrm{Cost}}(\mathrm{Cost})=1,\quad \mathrm{Cost}\le G_{\mathrm{Cost}}",
        r"\mu_{\mathrm{Cost}}(\mathrm{Cost})=\frac{C_{\max}-\mathrm{Cost}}{C_{\max}-G_{\mathrm{Cost}}},\quad G_{\mathrm{Cost}}<\mathrm{Cost}<C_{\max}",
        r"\mu_{\mathrm{Cost}}(\mathrm{Cost})=0,\quad \mathrm{Cost}\ge C_{\max}",
    ],
    eq_label="Equations (11a–11c). Cost membership (piecewise linear)"
)

add_par(
    "Equations (11a–11c) encode budget satisfaction: costs at or below the aspiration level G_Cost fully satisfy the cost goal, costs above C_max are unacceptable, "
    "and costs in between decrease satisfaction linearly."
)

add_par("Waiting membership function:")
add_equation_block(
    [
        r"\mu_{\mathrm{Wait}}(\mathrm{Wait})=1,\quad \mathrm{Wait}\le G_{\mathrm{Wait}}",
        r"\mu_{\mathrm{Wait}}(\mathrm{Wait})=\frac{W_{\max}-\mathrm{Wait}}{W_{\max}-G_{\mathrm{Wait}}},\quad G_{\mathrm{Wait}}<\mathrm{Wait}<W_{\max}",
        r"\mu_{\mathrm{Wait}}(\mathrm{Wait})=0,\quad \mathrm{Wait}\ge W_{\max}",
    ],
    eq_label="Equations (12a–12c). Waiting membership (piecewise linear)"
)

add_par(
    "Equations (12a–12c) define service satisfaction: waiting at or below G_Wait fully satisfies the service goal, waiting above W_max is unacceptable, "
    "and intermediate performance is graded linearly. In sensitivity studies, varying W_max traces feasibility boundaries and shows how tight tolerances "
    "restrict feasible operation under uncertainty."
)

add_par("Max–min MOFGPM objective (balanced tradeoff):")
add_equation_block(
    [
        r"\max\ \lambda",
        r"\lambda \le \mu_{\mathrm{Cost}},\quad \lambda \le \mu_{\mathrm{Wait}}",
        r"0 \le \lambda \le 1",
    ],
    eq_label="Equations (13–15). Max–min MOFGPM"
)

add_par(
    "Equations (13–15) enforce a balanced solution: λ improves only if both cost and waiting satisfaction improve, because λ is bounded by the smaller membership. "
    "This avoids solutions that are excellent on one dimension but unacceptable on the other and produces a managerial tradeoff map that is easy to interpret."
)

# ============================================================
# 6. SOLUTION APPROACH
# ============================================================
doc.add_heading("6. Solution approach: baseline and sensitivity pipeline", level=1)

add_par(
    "Waiting time depends on nonlinear interactions between arrivals, precedence, and shared resources. The pipeline therefore uses a simulation-optimization "
    "approach. For each staffing configuration, a genetic algorithm (GA) searches over task-priority permutations. Each chromosome is decoded into a feasible "
    "schedule using serial schedule generation under the precedence rules (Equations 4–6) and resource feasibility."
)

add_par(
    "The GA fitness is the MOFGPM objective: for each decoded schedule the pipeline computes Cost(S) (Equation 1), Wait(S) (Equation 9), memberships "
    "(Equations 11–12), and λ (Equations 13–15). Schedules violating triage feasibility (Equation 10) receive strong penalties so that feasibility is satisfied "
    "before improving satisfaction."
)

add_par(
    "Two execution modes align with baseline-anchored decision-making: (i) a baseline run with fixed staffing to quantify current feasibility and satisfaction, "
    "and (ii) a sensitivity mode that samples staffing vectors within bounds (Table 1) and/or varies waiting tolerance W_max to map feasibility boundaries and "
    "cost–service tradeoffs under optimistic/expected/pessimistic duration scenarios."
)

# ============================================================
# 7. FEASIBILITY BOUNDARIES + TRADEOFF REPORTING
# ============================================================
doc.add_heading("7. Feasibility boundaries and tradeoff reporting", level=1)

add_par(
    "Each tested staffing plan reports Cost, Wait, μ_Cost, μ_Wait, λ, and a feasibility flag (based primarily on triage feasibility and positive memberships). "
    "Aggregating results across staffing samples and uncertainty scenarios yields a decision-relevant map of feasible/infeasible regions and the achievable "
    "tradeoff frontier between cost and service satisfaction."
)

add_par(
    "For interpretation, λ provides a single summary index of balanced performance, while the individual memberships μ_Cost and μ_Wait identify which objective "
    "is binding. Reporting both λ and its components prevents tradeoffs from being hidden."
)

# ============================================================
# 8. VALIDATION
# ============================================================
doc.add_heading("8. Validation and robustness checks", level=1)

add_par(
    "Robustness is evaluated through repeated stochastic replications and confidence-interval summaries of waiting-time and triage-level service metrics. "
    "This ensures that baseline conclusions and sensitivity trends are not artifacts of a single simulated day."
)

doc.add_heading("8.1 External validation with FlexSim", level=2)
add_par(
    "In addition to the methodology runs (baseline and sensitivity), a FlexSim discrete-event simulation can be used to validate the GA-based approach. "
    "Baseline staffing and selected sensitivity staffing samples are replicated in FlexSim using the same triage pathways, arrival logic, and service assumptions."
)
add_par(
    "Validation compares key outputs between GA+MOFGPM and FlexSim, including total waiting time, triage-level service attainment, and resource utilization. "
    "Close agreement supports the validity of the optimization pipeline and its managerial interpretations."
)


def build_overleaf_tex():
    return r"""\documentclass[12pt]{article}
\usepackage[a4paper,margin=1in]{geometry}
\usepackage{setspace}
\usepackage{amsmath,amssymb,mathtools}
\usepackage{booktabs}
\usepackage{newtxtext,newtxmath}
\setstretch{1.15}

\title{Methodology: Baseline-Anchored Sensitivity Analysis with MOFGPM and Jim{\'e}nez Linearization}
\date{}

\begin{document}
\maketitle

This section presents a decision-support pipeline for Emergency Department (ED) resource planning under uncertainty. The methodology is explicitly \emph{baseline-anchored}: it begins from a realistic staffing configuration already used (or constrained) in practice, and then extends the analysis through systematic sensitivity runs to (i) diagnose feasibility of service targets under uncertainty and (ii) map interpretable tradeoffs between service satisfaction and cost. Rather than producing a single ``optimal'' plan, the pipeline produces a feasibility envelope and a cost--service tradeoff surface that managers can use for negotiation and planning.

\section{Problem statement and baseline-anchored design}
Patients arrive over time, are assigned a triage category, and follow triage-dependent activity routes. Activities consume shared resources (Doctor, Nurse, Assistant, Specialist). A staffing plan specifies available capacity per resource type. Given a staffing plan, we construct a feasible schedule that respects patient precedence and resource availability, compute waiting-time performance and staffing cost, and quantify goal satisfaction using max--min Multiobjective Fuzzy Goal Programming (MOFGPM).

The baseline-anchored design directly addresses the managerial gap identified in the Introduction: decision-makers do not start from scratch but from an existing staffing configuration. Consequently, the primary outputs are (a) baseline feasibility and satisfaction, (b) feasibility boundaries under uncertainty, and (c) cost--service tradeoffs for alternative staffing levels.

\section{Case inputs: resources, staffing, and cost}
Staffing is analyzed in two modes: (i) a baseline fixed staffing configuration representing current operations or policy constraints, and (ii) sensitivity sampling within lower/upper bounds to explore alternative capacity levels. Table~\ref{tab:staff} summarizes resource types, unit costs, baseline staffing, and practical bounds used in sensitivity runs.

\begin{table}[h!]
\centering
\begin{tabular}{lcccc}
\toprule
Resource type & Unit cost & Baseline staff & Lower bound & Upper bound \\
\midrule
Doctor & 200 & 3 & 1 & 6 \\
Nurse & 100 & 3 & 1 & 8 \\
Assistant & 50 & 6 & 1 & 10 \\
Specialist & 300 & 1 & 0 & 2 \\
\bottomrule
\end{tabular}
\caption{Resource types, unit costs, baseline staffing, and sensitivity bounds.}
\label{tab:staff}
\end{table}

For a staffing vector $S=\{S_r\}_{r\in\mathcal{R}}$, the total staffing cost is the linear aggregation of unit costs and staffing quantities:
\begin{equation}
\mathrm{Cost}(S) = \sum_{r\in\mathcal{R}} SC_{r}\,S_{r}.
\label{eq:cost}
\end{equation}
Equation~\eqref{eq:cost} links \emph{capacity decisions} to \emph{financial sustainability}: increasing staffing increases the capacity available to meet demand but also increases cost. This cost term later enters MOFGPM through a satisfaction (membership) function that reflects a manager’s budget aspiration and tolerance.

\section{Uncertainty modeling}
\subsection{Patient arrivals by triage (demand model)}
Arrival streams are generated by fitted distributions per triage level, then accumulated into a representative day timeline. Let $TA_i$ denote the arrival time of patient $i$ and $c(i)$ denote the triage class of patient $i$:
\begin{equation}
TA_{i}: \text{arrival time of patient } i, \qquad c(i): \text{triage class of patient } i.
\label{eq:arrivals}
\end{equation}
Equation~\eqref{eq:arrivals} defines the \emph{demand side} of the system. The collection of arrival times $\{TA_i\}$ determines when patients enter the system and therefore how demand overlaps over time. Holding the same realized arrival stream across uncertainty scenarios (when desired) isolates the effect of service-time uncertainty and staffing on feasibility and performance.

\subsection{Service-time uncertainty and fuzzy mapping}
Service durations are uncertain due to heterogeneity in patient condition, operational interruptions, and clinical complexity. To integrate heterogeneous uncertainty descriptions into one framework, we map each activity duration to a fuzzy triple $(p,m,o)$ representing an optimistic value $p$, a central value $m$, and a pessimistic value $o$. This supports scenario-based analysis while maintaining a consistent defuzzification rule.

For triangular activities, we use $(p,m,o)=(L,M,U)$ where $M$ is the most likely duration. For uniform activities over $(a,b)$ we use $(p,m,o)=\left(a,\frac{a+b}{2},b\right)$. For exponential activities, we retain Jim{\'e}nez’s linearization to produce a bounded representation from a location parameter $a$ and scale $\sigma$:
\begin{align}
(p,m,o) &= (L,M,U) && \text{triangular} \label{eq:fuzzy_tri}\\
(p,m,o) &= \left(a,\frac{a+b}{2},b\right) && \text{uniform}(a,b) \label{eq:fuzzy_unif}\\
(p,m,o) &= \left(a,\;a+\frac{\sigma}{2},\;a+\sigma\right) && \text{exponential via Jim{\'e}nez}. \label{eq:fuzzy_jimenez}
\end{align}
Equations~\eqref{eq:fuzzy_tri}--\eqref{eq:fuzzy_jimenez} define how uncertainty in durations is represented. In particular, Equation~\eqref{eq:fuzzy_jimenez} preserves the Jim{\'e}nez interval concept $[a,a+\sigma]$ while defining a central value $a+\sigma/2$ to create a triangular-like fuzzy triple. In implementation, $a$ may be set to the mean or median of the fitted exponential distribution depending on robustness needs.

To evaluate schedules, the pipeline converts the fuzzy triple to a single crisp duration using an expected-value defuzzification operator:
\begin{equation}
EI(\tilde{T}) = \frac{p + 2m + o}{4}.
\label{eq:defuzz}
\end{equation}
Equation~\eqref{eq:defuzz} provides the \emph{expected scenario} duration used in the scheduling module. Conceptually, it balances optimistic and pessimistic extremes while placing higher weight on the central value, ensuring that the schedule evaluation reflects both uncertainty and clinical plausibility.

\section{Scheduling-based evaluation (capacity meets demand)}
Given arrivals (demand) and staffing (capacity), the pipeline constructs a feasible schedule for all required patient activities. Let $ST_{ij}$ and $ET_{ij}$ denote start and end times of activity $j$ for patient $i$. Let $d_{ij}$ be the defuzzified duration, typically $d_{ij}=EI(\tilde{T}_{ij})$ from Equation~\eqref{eq:defuzz}. Precedence and timing are:
\begin{align}
ST_{i1} &\ge TA_{i} \label{eq:prec1}\\
ST_{ij} &\ge ET_{i,j-1}, \quad \forall j>1 \label{eq:prec2}\\
ET_{ij} &= ST_{ij} + d_{ij}. \label{eq:endtime}
\end{align}
Equations~\eqref{eq:prec1}--\eqref{eq:endtime} ensure clinical and temporal feasibility. Equation~\eqref{eq:prec1} prevents service before arrival (no demand before entry). Equation~\eqref{eq:prec2} enforces within-patient precedence (the route must be followed). Equation~\eqref{eq:endtime} links start times to completion times using uncertainty-adjusted durations.

Resource feasibility (the \emph{capacity side}) is enforced through a serial schedule generation scheme: when a task requires one or more resource types, it is assigned to the earliest available units of each required resource type, and its start time is set to the latest of (i) the precedence-implied release time and (ii) the selected resource availability times. This operationally enforces the ``balance between demand and capacity'': if arrivals overlap heavily or staffing is low, resource availability delays start times and increases waiting.

\subsection{Waiting-time metric and triage feasibility}
Waiting time is computed as the time a patient spends waiting between arrival or completion of a prior activity and the start of the next activity:
\begin{align}
WT_{i1} &= ST_{i1} - TA_{i} \label{eq:wt1}\\
WT_{ij} &= ST_{ij} - ET_{i,j-1}, \quad \forall j>1 \label{eq:wt2}\\
\mathrm{Wait}(S) &= \sum_{i}\sum_{j} WT_{ij}. \label{eq:wait_total}
\end{align}
Equations~\eqref{eq:wt1}--\eqref{eq:wait_total} define the service-quality metric used in MOFGPM. The aggregate $\mathrm{Wait}(S)$ is interpretable as total accumulated waiting time across all patients and activities under staffing plan $S$. This quantity increases when demand intensity (arrivals) or service durations exceed available capacity, and it decreases when additional staffing increases capacity.

In addition to aggregate waiting, triage categories impose service-level requirements. Let $WT^{\mathrm{early}}_i$ denote cumulative waiting time over early-care activities (e.g., up to $j\le J_0$ with $J_0=4$ in the implementation), and let $MT_{c}$ denote the triage-specific maximum tolerated threshold. Feasibility requires:
\begin{equation}
WT^{\mathrm{early}}_i \le MT_{c(i)}, \quad \forall i.
\label{eq:triage_sla}
\end{equation}
Equation~\eqref{eq:triage_sla} provides a manager- and clinician-relevant feasibility rule: even if the aggregate waiting time is acceptable on average, excessive early-care delays for high-acuity patients are not permitted. In the computational pipeline, violations of \eqref{eq:triage_sla} are heavily penalized so that solutions are first driven toward feasibility before improving satisfaction.

\section{Multiobjective Fuzzy Goal Programming (MOFGPM)}
Cost and waiting are conflicting objectives: increasing staffing reduces waiting but increases cost. Instead of using fixed weights (which are difficult to justify operationally), MOFGPM converts each objective into a satisfaction degree in $[0,1]$ using piecewise-linear membership functions defined by aspiration levels (Goal) and unacceptable limits (Max).

For cost:
\begin{align}
\mu_{\mathrm{Cost}}(\mathrm{Cost}) &= 1, && \mathrm{Cost} \le G_{\mathrm{Cost}} \label{eq:mu_cost_1}\\
\mu_{\mathrm{Cost}}(\mathrm{Cost}) &= \frac{C_{\max}-\mathrm{Cost}}{C_{\max}-G_{\mathrm{Cost}}}, && G_{\mathrm{Cost}}<\mathrm{Cost}<C_{\max} \label{eq:mu_cost_2}\\
\mu_{\mathrm{Cost}}(\mathrm{Cost}) &= 0, && \mathrm{Cost} \ge C_{\max}. \label{eq:mu_cost_3}
\end{align}
Equations~\eqref{eq:mu_cost_1}--\eqref{eq:mu_cost_3} encode the idea that costs at or below $G_{\mathrm{Cost}}$ fully satisfy the budget goal, costs above $C_{\max}$ are unacceptable, and costs in between reduce satisfaction linearly.

For waiting:
\begin{align}
\mu_{\mathrm{Wait}}(\mathrm{Wait}) &= 1, && \mathrm{Wait} \le G_{\mathrm{Wait}} \label{eq:mu_wait_1}\\
\mu_{\mathrm{Wait}}(\mathrm{Wait}) &= \frac{W_{\max}-\mathrm{Wait}}{W_{\max}-G_{\mathrm{Wait}}}, && G_{\mathrm{Wait}}<\mathrm{Wait}<W_{\max} \label{eq:mu_wait_2}\\
\mu_{\mathrm{Wait}}(\mathrm{Wait}) &= 0, && \mathrm{Wait} \ge W_{\max}. \label{eq:mu_wait_3}
\end{align}
Equations~\eqref{eq:mu_wait_1}--\eqref{eq:mu_wait_3} define a service-quality satisfaction scale: the ED is fully satisfactory if the waiting metric is at or below $G_{\mathrm{Wait}}$, unacceptable if it exceeds $W_{\max}$, and graded linearly in between. Importantly, in sensitivity studies $W_{\max}$ can be varied to trace feasibility boundaries and to show how ``tight'' waiting tolerances restrict feasible operation under uncertainty.

The max--min MOFGPM objective then maximizes the minimum satisfaction level $\lambda$ across objectives:
\begin{align}
\max \ & \lambda \label{eq:mofgpm_obj}\\
\text{s.t. } & \lambda \le \mu_{\mathrm{Cost}}, \quad \lambda \le \mu_{\mathrm{Wait}} \label{eq:mofgpm_cons}\\
& 0 \le \lambda \le 1. \label{eq:mofgpm_bounds}
\end{align}
Equations~\eqref{eq:mofgpm_obj}--\eqref{eq:mofgpm_bounds} enforce a \emph{balanced} tradeoff: $\lambda$ improves only if both cost and waiting satisfaction improve, since $\lambda$ is bounded by the smaller membership. This directly supports the managerial goal of avoiding solutions that are excellent on one dimension but unacceptable on the other.

\section{Solution approach: baseline and sensitivity pipeline}
The optimization stage uses a genetic algorithm (GA) that searches over task-priority permutations and decodes each chromosome into a feasible schedule under precedence and resource feasibility rules. The GA fitness is defined by the MOFGPM objective: for each decoded schedule, we compute $\mathrm{Cost}(S)$ via Equation~\eqref{eq:cost}, $\mathrm{Wait}(S)$ via Equation~\eqref{eq:wait_total}, memberships via Equations~\eqref{eq:mu_cost_1}--\eqref{eq:mu_wait_3}, and $\lambda$ via Equations~\eqref{eq:mofgpm_obj}--\eqref{eq:mofgpm_bounds}. Schedules violating the triage feasibility rule in Equation~\eqref{eq:triage_sla} receive strong penalties.

Two execution modes are used:
(i) \textbf{baseline case} with fixed staffing to quantify current feasibility and satisfaction, and
(ii) \textbf{sensitivity analysis} by sampling staffing vectors within practical bounds (Table~\ref{tab:staff}) and/or varying $W_{\max}$ to map feasibility boundaries and cost--service tradeoffs under optimistic/expected/pessimistic duration scenarios.

\section{Feasibility boundaries and tradeoff reporting}
Each tested staffing plan reports $\mathrm{Cost}$, $\mathrm{Wait}$, $\mu_{\mathrm{Cost}}$, $\mu_{\mathrm{Wait}}$, $\lambda$, and a feasibility flag determined primarily by Equation~\eqref{eq:triage_sla} and by whether both memberships are positive (i.e., performance remains within tolerance limits). Aggregating these outputs produces (a) feasible/infeasible regions and (b) an achievable tradeoff frontier between cost and service satisfaction.

\section{Validation and robustness checks}
Robustness is evaluated through repeated stochastic replications and confidence-interval summaries of waiting-time and triage-level service metrics. This step ensures that baseline conclusions and sensitivity trends are not artifacts of a single simulated day.

\subsection{External validation with FlexSim}
In addition to the methodology runs (baseline and sensitivity), a FlexSim discrete-event simulation can be used to validate the GA-based approach. Baseline staffing and selected sensitivity staffing samples are replicated in FlexSim using the same triage pathways, arrival logic, and service assumptions. Validation compares key outputs between GA+MOFGPM and FlexSim (e.g., total waiting time, triage-level service attainment, and resource utilization). Close agreement supports the validity of the optimization pipeline and its managerial interpretations.

\end{document}
"""


tex_path = Path(__file__).resolve().parent / "Methodology_MOFGPM_Pipeline.tex"
tex_path.write_text(build_overleaf_tex(), encoding="utf-8")
print("Saved:", tex_path)


def compile_latex_to_pdf(tex_file: Path):
    pdflatex = shutil.which("pdflatex")
    if not pdflatex:
        print("pdflatex not found on PATH. Skipping PDF generation.")
        return

    workdir = tex_file.parent
    cmd = [pdflatex, "-interaction=nonstopmode", "-halt-on-error", tex_file.name]
    try:
        # Run twice to stabilize references/equation numbering if needed.
        for _ in range(2):
            result = subprocess.run(cmd, cwd=workdir, capture_output=True, text=True)
            if result.returncode != 0:
                print("PDF generation failed.")
                print(result.stdout[-1200:])
                print(result.stderr[-1200:])
                return
        print("Saved:", tex_file.with_suffix(".pdf"))
    except Exception as exc:
        print("PDF generation error:", exc)


compile_latex_to_pdf(tex_path)

out_path = Path(__file__).resolve().parent / "Methodology_MOFGPM_Pipeline.docx"
try:
    doc.save(out_path)
    print("Saved:", out_path)
except PermissionError:
    print("Could not save DOCX (file is open/locked):", out_path)
