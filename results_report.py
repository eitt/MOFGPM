def generate_results_report_word(
    records_df,
    out_path,
    *,
    title="Results - Baseline-Anchored Feasibility and Tradeoff Analysis (MOFGPM)",
    baseline_staff=None,
    goals=None,
    wait_metric_desc="Total accumulated waiting time across all patients and activities (minutes).",
    feasibility_rule_desc="Feasible if triage early-care standards are satisfied and both memberships are positive (mu_Cost>0, mu_Wait>0).",
    figures=None,
    scenario_col="scenario",
    wmax_col="mw",
):
    import numpy as np
    import pandas as pd
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    df = records_df.copy()
    if not isinstance(df, pd.DataFrame):
        raise TypeError("records_df must be a pandas DataFrame")

    required = {scenario_col, "cost", "wait", "mu_cost", "mu_wait", "lambda", "feasible"}
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise ValueError(f"records_df is missing required columns: {missing}")

    df["feasible"] = df["feasible"].astype(bool)
    for c in ["cost", "wait", "mu_cost", "mu_wait", "lambda"]:
        df[c] = pd.to_numeric(df[c], errors="coerce")

    candidate_resources = ["Doctor", "Nurse", "Assistant", "Specialist"]
    resource_cols = [c for c in candidate_resources if c in df.columns]

    has_wmax = wmax_col in df.columns
    if has_wmax:
        df[wmax_col] = pd.to_numeric(df[wmax_col], errors="coerce")

    def safe_float(x):
        try:
            return float(x)
        except Exception:
            return np.nan

    def pct(a, b):
        return 0.0 if b == 0 else 100.0 * a / b

    def best_feasible(sub):
        s = sub[sub["feasible"] & np.isfinite(sub["lambda"])]
        if s.empty:
            return None
        s = s.sort_values(["lambda", "cost", "wait"], ascending=[False, True, True])
        return s.iloc[0]

    def min_wmax_feasible(sub):
        if not has_wmax:
            return np.nan
        s = sub[sub["feasible"] & np.isfinite(sub[wmax_col])]
        if s.empty:
            return np.nan
        return float(np.min(s[wmax_col]))

    def binding(mu_c, mu_w):
        if not np.isfinite(mu_c) or not np.isfinite(mu_w):
            return "Unverified"
        if abs(mu_c - mu_w) < 1e-6:
            return "Balanced"
        return "Cost" if mu_c < mu_w else "Wait"

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for h in ["Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[h]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    def add_par(text, bold=False, italic=False, align=None, space_after=6, space_before=0):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if align is not None:
            p.alignment = align
        pf = p.paragraph_format
        pf.space_after = Pt(space_after)
        pf.space_before = Pt(space_before)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        return p

    def add_caption(text):
        add_par(text, italic=True, space_before=2, space_after=10)

    def add_table(df_table, caption=None):
        t = doc.add_table(rows=1, cols=len(df_table.columns))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for j, c in enumerate(df_table.columns):
            hdr[j].text = str(c)
        for _, row in df_table.iterrows():
            cells = t.add_row().cells
            for j, c in enumerate(df_table.columns):
                cells[j].text = str(row[c])
        if caption:
            add_caption(caption)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)

    add_par(
        "This Results section is structured to match the baseline-anchored decision gap described in the Introduction. "
        "We first report baseline performance, then diagnose feasibility boundaries under uncertainty, and finally map the "
        "cost-service tradeoff using MOFGPM outcomes (mu_Cost, mu_Wait, and lambda)."
    )

    doc.add_heading("1. Inputs and Evaluation Definitions", level=1)
    if baseline_staff:
        add_par("Baseline staffing configuration (reference case):", bold=True, space_after=2)
        add_par(", ".join([f"{k}={v}" for k, v in baseline_staff.items()]), space_after=8)
    if goals:
        add_par("MOFGPM goal and tolerance parameters:", bold=True, space_after=2)
        goal_lines = []
        for k in ["G_cost", "C_max", "G_wait", "W_max"]:
            if k in goals:
                goal_lines.append(f"{k}={goals[k]}")
        add_par(", ".join(goal_lines) if goal_lines else "[Unverified]", space_after=8)
    add_par(f"Waiting metric used: {wait_metric_desc}")
    add_par(f"Feasibility rule: {feasibility_rule_desc}", space_after=10)

    doc.add_heading("2. Baseline Case Results", level=1)
    baseline_row = None
    if baseline_staff and resource_cols:
        mask = np.ones(len(df), dtype=bool)
        for r in resource_cols:
            if r in baseline_staff:
                mask &= (df[r] == baseline_staff[r])
        cand = df[mask].copy()
        if not cand.empty:
            bf = best_feasible(cand)
            baseline_row = bf if bf is not None else cand.iloc[0]

    if baseline_row is None:
        add_par(
            "A unique baseline row could not be identified from the provided results table. "
            "If you want an explicit baseline subsection, store the baseline staffing vector in the results DataFrame."
        )
    else:
        feas_txt = "Feasible" if bool(baseline_row["feasible"]) else "Infeasible"
        add_par(
            f"Under baseline staffing, the system is {feas_txt}. Cost={safe_float(baseline_row['cost']):.2f}, "
            f"Wait={safe_float(baseline_row['wait']):.2f}, mu_Cost={safe_float(baseline_row['mu_cost']):.3f}, "
            f"mu_Wait={safe_float(baseline_row['mu_wait']):.3f}, lambda={safe_float(baseline_row['lambda']):.3f}."
        )
        add_par(
            f"The binding objective at baseline is {binding(baseline_row['mu_cost'], baseline_row['mu_wait'])}."
        )

    doc.add_heading("3. Feasibility Under Uncertainty Scenarios", level=1)
    scenarios = sorted(list(df[scenario_col].dropna().unique()))
    scenario_label = "Operational scenario" if scenario_col == "scenario_key" else "Scenario"
    is_merged = ("run_tag" in df.columns and df["run_tag"].nunique(dropna=True) > 1) or (
        scenario_col == "scenario_key"
    )
    if is_merged:
        n_runs = int(df["run_tag"].nunique()) if "run_tag" in df.columns else 1
        n_keys = int(df[scenario_col].nunique(dropna=True))
        add_par(
            f"This section integrates merged outputs across {n_runs} run batches and {n_keys} operational scenario groups."
        )

    rows = []
    for sc in scenarios:
        sub = df[df[scenario_col] == sc]
        n = len(sub)
        n_feas = int(sub["feasible"].sum())
        feas_rate = pct(n_feas, n)
        bf = best_feasible(sub)
        min_w = min_wmax_feasible(sub)
        if bf is None:
            rows.append(
                {
                    scenario_label: sc,
                    "Runs": n,
                    "Feasible (%)": f"{feas_rate:.1f}",
                    "Min feasible W_max": f"{min_w:.2f}" if np.isfinite(min_w) else "-",
                    "Best lambda": "-",
                    "Best Cost": "-",
                    "Best Wait": "-",
                    "Binding": "-",
                }
            )
        else:
            rows.append(
                {
                    scenario_label: sc,
                    "Runs": n,
                    "Feasible (%)": f"{feas_rate:.1f}",
                    "Min feasible W_max": f"{min_w:.2f}" if np.isfinite(min_w) else "-",
                    "Best lambda": f"{safe_float(bf['lambda']):.3f}",
                    "Best Cost": f"{safe_float(bf['cost']):.2f}",
                    "Best Wait": f"{safe_float(bf['wait']):.2f}",
                    "Binding": binding(bf["mu_cost"], bf["mu_wait"]),
                }
            )

    add_par("Table 1 summarizes feasibility rates and best-achieved balanced satisfaction by analysis group.")
    add_table(pd.DataFrame(rows), caption="Table 1. Scenario-level feasibility and best MOFGPM outcomes.")

    doc.add_heading("4. Cross-Scenario Integrated Diagnostics", level=1)
    integ_rows = []
    for sc in scenarios:
        sub = df[df[scenario_col] == sc]
        feas = sub[sub["feasible"] & np.isfinite(sub["lambda"])]
        integ_rows.append(
            {
                scenario_label: sc,
                "Samples": len(sub),
                "Feasible (%)": f"{pct(int(sub['feasible'].sum()), len(sub)):.1f}",
                "Median lambda (feasible)": f"{float(feas['lambda'].median()):.3f}" if not feas.empty else "-",
                "P90 lambda (feasible)": f"{float(feas['lambda'].quantile(0.90)):.3f}" if not feas.empty else "-",
                "Mean cost": f"{float(pd.to_numeric(sub['cost'], errors='coerce').mean()):.2f}",
                "Mean wait": f"{float(pd.to_numeric(sub['wait'], errors='coerce').mean()):.2f}",
            }
        )
    add_table(
        pd.DataFrame(integ_rows),
        caption="Table 2. Integrated performance diagnostics across scenario groups.",
    )

    doc.add_heading("5. Cost-Service Tradeoff Map (MOFGPM Interpretation)", level=1)
    feasible = df[df["feasible"] & np.isfinite(df["lambda"])].copy()
    if feasible.empty:
        add_par("No feasible solutions were found in the provided results.")
    else:
        topk = feasible.sort_values(["lambda", "cost", "wait"], ascending=[False, True, True]).head(10)
        cols = [scenario_col]
        if has_wmax:
            cols.append(wmax_col)
        cols += resource_cols + ["cost", "wait", "mu_cost", "mu_wait", "lambda"]
        cols = [c for c in cols if c in topk.columns]
        topk_disp = topk[cols].copy()
        for c in ["cost", "wait"]:
            if c in topk_disp.columns:
                topk_disp[c] = topk_disp[c].map(lambda x: f"{float(x):.2f}" if np.isfinite(x) else "-")
        for c in ["mu_cost", "mu_wait", "lambda"]:
            if c in topk_disp.columns:
                topk_disp[c] = topk_disp[c].map(lambda x: f"{float(x):.3f}" if np.isfinite(x) else "-")
        if has_wmax and wmax_col in topk_disp.columns:
            topk_disp[wmax_col] = topk_disp[wmax_col].map(
                lambda x: f"{float(x):.2f}" if np.isfinite(x) else "-"
            )
        add_table(topk_disp, caption="Table 3. Top feasible solutions (ranked by lambda, then cost and wait).")

    doc.add_heading("6. Feasibility Boundary Diagnostics", level=1)
    if not has_wmax:
        add_par(
            f"A W_max sweep was not detected (column '{wmax_col}' not found)."
        )
    else:
        boundary_rows = []
        for sc in scenarios:
            sub = df[df[scenario_col] == sc]
            mfw = min_wmax_feasible(sub)
            boundary_rows.append(
                {
                    scenario_label: sc,
                    "Min feasible W_max": f"{mfw:.2f}" if np.isfinite(mfw) else "-",
                }
            )
        add_table(pd.DataFrame(boundary_rows), caption="Table 4. Minimum feasible W_max by scenario group.")

    if figures:
        doc.add_heading("7. Figures", level=1)
        for cap, path in figures.items():
            try:
                doc.add_picture(path, width=Inches(6.5))
                add_caption(cap)
            except Exception:
                add_par(f"[Unverified] Could not insert figure: {cap} ({path})")

    doc.add_heading("8. Summary", level=1)
    add_par(
        "Results are presented as baseline performance plus integrated feasibility and tradeoff diagnostics. "
        "Merged exploration outputs (random sampling and combinatorics) identify robust feasible regions, likely binding objectives, and the cost associated with service improvements."
    )

    doc.save(out_path)
    return out_path
